#!/usr/bin/env python3
"""Convert PDF to DOCX with text + images preserved."""
import sys
import os
from pathlib import Path
import io

def convert_pdf_to_docx(pdf_path, docx_path):
    try:
        import fitz  # PyMuPDF
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError as e:
        print(f"Missing dependency: {e}")
        print("Install with: pip install PyMuPDF python-docx")
        sys.exit(1)

    pdf_path = Path(pdf_path)
    docx_path = Path(docx_path)

    if not pdf_path.exists():
        print(f"PDF not found: {pdf_path}")
        sys.exit(1)

    doc = Document()

    # Set page margins (0.5 inch for more content)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    docx_path.parent.mkdir(parents=True, exist_ok=True)

    pdf_doc = fitz.open(str(pdf_path))
    total_pages = len(pdf_doc)
    print(f"Converting {total_pages} pages...")

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("Appian ASD STIG V6R4 AdminConsole Evidence")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Application Security and Development STIG Version: 6, Release: 4")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing

    temp_dir = docx_path.parent / "_temp_images"
    temp_dir.mkdir(exist_ok=True)

    try:
        for page_num in range(total_pages):
            page = pdf_doc[page_num]
            print(f"  Processing page {page_num + 1}/{total_pages}...")

            # Extract text
            text = page.get_text("text").strip()

            if text:
                # Add page separator
                if page_num > 0:
                    doc.add_paragraph()
                    sep = doc.add_paragraph()
                    sep_run = sep.add_run("─" * 80)
                    sep_run.font.size = Pt(6)
                    sep_run.font.color.rgb = RGBColor(200, 200, 200)
                    doc.add_paragraph()

                # Split text into lines and add
                lines = text.split("\n")
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue

                    para = doc.add_paragraph()
                    run = para.add_run(line)

                    # Detect headers (V- numbers, bold them)
                    if line.startswith("V-") and any(c.isdigit() for c in line):
                        run.bold = True
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 51, 102)
                    elif "[Evidence Screenshot Placeholder]" in line:
                        run.italic = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(100, 100, 100)
                    elif line.startswith("Explanation/Context:"):
                        run.bold = True
                        run.font.size = Pt(10)
                    else:
                        run.font.size = Pt(10)

            # Extract images from the page
            images = page.get_images(full=True)
            for img_index, img_info in enumerate(images):
                xref = img_info[0]
                try:
                    pix = fitz.Pixmap(pdf_doc, xref)
                    if pix.width < 50 or pix.height < 50:  # Skip tiny images
                        pix = None
                        continue

                    # Convert to RGB if needed
                    if pix.n > 4:
                        pix = fitz.Pixmap(fitz.csRGB, pix)

                    img_path = temp_dir / f"page{page_num + 1}_img{img_index}.png"
                    pix.save(str(img_path))
                    pix = None

                    if img_path.exists() and img_path.stat().st_size > 1024:
                        doc.add_picture(str(img_path), width=Inches(6.0))
                        last_para = doc.paragraphs[-1]
                        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph()  # spacing after image

                except Exception as e:
                    print(f"    Warning: Could not extract image {img_index} on page {page_num + 1}: {e}")
                    continue

    finally:
        pdf_doc.close()

        # Clean up temp images
        for f in temp_dir.glob("*.png"):
            try:
                f.unlink()
            except:
                pass
        try:
            temp_dir.rmdir()
        except:
            pass

    # Save the document
    doc.save(str(docx_path))
    print(f"\nSaved to: {docx_path}")
    print(f"File size: {docx_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python pdf_to_docx.py <input.pdf> [output.docx]")
        sys.exit(1)

    pdf_path = sys.argv[1]
    if len(sys.argv) >= 3:
        docx_path = sys.argv[2]
    else:
        docx_path = str(Path(pdf_path).with_suffix(".docx"))

    convert_pdf_to_docx(pdf_path, docx_path)
