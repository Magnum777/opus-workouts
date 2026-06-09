import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

# Load all Not a Finding items from the new CSV
path = r"C:\Users\compj\.openclaw\media\inbound\STIG_APPIAN_REVIEWED---530feffe-8008-47bf-9dae-3282b892629d.csv"
items = {}
with open(path, 'r', encoding='utf-8', newline='') as f:
    f.readline()
    for row in csv.DictReader(f):
        status = (row.get('Status') or '').strip().lower()
        if status == 'not a finding':
            gid = row['Group ID'].strip()
            items[gid] = row

print(f"Loaded {len(items)} Not a Finding items")

# ============================================================
# Build comprehensive evidence package per PIEE guide
# ============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# --- A. TITLE ---
title = doc.add_paragraph()
title_run = title.add_run("Application Security and Development STIG\nVersion: 6, Release: 4")
title_run.font.size = Pt(18)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

cui = doc.add_paragraph()
cui_run = cui.add_run("UNCLASSIFIED//CUI")
cui_run.font.size = Pt(10)
cui_run.font.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run(f"Appian Low-Code Platform — Evidence Package\n{len(items)} Not a Finding Items | Date: {datetime.now().strftime('%Y-%m-%d')}")
subtitle_run.font.size = Pt(11)
subtitle_run.font.italic = True

doc.add_paragraph()

# --- B. TABLE OF CONTENTS (grouped by category) ---
contents_h = doc.add_paragraph()
contents_h_run = contents_h.add_run("Contents")
contents_h_run.font.size = Pt(14)
contents_h_run.font.bold = True
contents_h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

# Group by category/theme
categories = {
    "Session Management": [],
    "Authentication & Password": [],
    "Audit & Logging": [],
    "Access Control": [],
    "Data Protection": [],
    "Security Attributes / Marking": [],
    "Cryptography / PKI / SAML": [],
    "Application Security": [],
    "Configuration & Maintenance": [],
    "Other": []
}

for gid, item in items.items():
    rt = (item.get('Rule Title') or '').lower()
    if any(k in rt for k in ['session', 'idle', 'timeout', 'logoff', 'cookie']):
        categories["Session Management"].append(item)
    elif any(k in rt for k in ['password', 'authenticat', 'logon', 'login', 'account', 'lockout', 'pki', 'cac', 'certificate']):
        categories["Authentication & Password"].append(item)
    elif any(k in rt for k in ['audit', 'log', 'record']):
        categories["Audit & Logging"].append(item)
    elif any(k in rt for k in ['access', 'privilege', 'authoriz', 'permission', 'role']):
        categories["Access Control"].append(item)
    elif any(k in rt for k in ['encrypt', 'cryptograph', 'data protection', 'fips', 'tls', 'ssl']):
        categories["Data Protection"].append(item)
    elif any(k in rt for k in ['mark', 'classif', 'banner', 'attribute', 'cui']):
        categories["Security Attributes / Marking"].append(item)
    elif any(k in rt for k in ['saml', 'ws-security', 'soap', 'assertion', 'signature']):
        categories["Cryptography / PKI / SAML"].append(item)
    elif any(k in rt for k in ['cross-site', 'sql injection', 'xss', 'csrf', 'input', 'validat', 'error', 'denial of service', 'dos']):
        categories["Application Security"].append(item)
    elif any(k in rt for k in ['patch', 'update', 'back-up', 'backup', 'decommission', 'configur']):
        categories["Configuration & Maintenance"].append(item)
    else:
        categories["Other"].append(item)

page_num = 3
for cat_name, cat_items in categories.items():
    if not cat_items:
        continue
    cat_line = doc.add_paragraph()
    cat_run = cat_line.add_run(f"{cat_name} ({len(cat_items)} items)")
    cat_run.font.size = Pt(11)
    cat_run.font.bold = True
    cat_line.paragraph_format.space_after = Pt(2)
    
    for item in cat_items:
        gid = item['Group ID']
        sid = item['STIG ID']
        toc_entry = doc.add_paragraph()
        toc_entry_run = toc_entry.add_run(f"  {gid} | {sid}")
        toc_entry_run.font.size = Pt(9)
        toc_entry.paragraph_format.space_after = Pt(1)
        page_num += 1

doc.add_page_break()

# ============================================================
# EVIDENCE BY CATEGORY
# ============================================================
for cat_name, cat_items in categories.items():
    if not cat_items:
        continue
    
    # Category header
    cat_h = doc.add_paragraph()
    cat_h_run = cat_h.add_run(f"=== {cat_name} ===")
    cat_h_run.font.size = Pt(14)
    cat_h_run.font.bold = True
    cat_h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    cat_h.paragraph_format.space_after = Pt(8)
    
    for item in cat_items:
        gid = item['Group ID']
        sid = item['STIG ID']
        sev = (item['Severity'] or '').upper()
        title_text = item['Rule Title']
        comments = item.get('Comments', '').strip()
        finding_details = item.get('Finding Details', '').strip()
        
        # Compact Vuln ID header
        h = doc.add_paragraph()
        h_run = h.add_run(f"{gid} | {sid} | {sev}")
        h_run.font.size = Pt(10)
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)
        h.paragraph_format.space_after = Pt(2)
        
        # Rule Title
        rt = doc.add_paragraph()
        rt_run = rt.add_run(title_text)
        rt_run.font.size = Pt(9)
        rt.paragraph_format.space_after = Pt(4)
        
        # Status
        status_p = doc.add_paragraph()
        status_label = status_p.add_run("Status: ")
        status_label.font.size = Pt(9)
        status_label.font.bold = True
        status_val = status_p.add_run("Not a Finding")
        status_val.font.size = Pt(9)
        status_val.font.bold = True
        status_val.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
        status_p.paragraph_format.space_after = Pt(2)
        
        # Finding Details (from CSV or generated)
        if finding_details:
            fd_p = doc.add_paragraph()
            fd_label = fd_p.add_run("Finding Details: ")
            fd_label.font.size = Pt(9)
            fd_label.font.bold = True
            fd_text = fd_p.add_run(finding_details[:300])
            fd_text.font.size = Pt(9)
            fd_p.paragraph_format.line_spacing = 1.1
            fd_p.paragraph_format.space_after = Pt(2)
        
        # Comments
        if comments:
            comm_p = doc.add_paragraph()
            comm_label = comm_p.add_run("Comments: ")
            comm_label.font.size = Pt(9)
            comm_label.font.bold = True
            comm_text = comm_p.add_run(comments[:400])
            comm_text.font.size = Pt(9)
            comm_p.paragraph_format.line_spacing = 1.1
            comm_p.paragraph_format.space_after = Pt(6)
        
        doc.add_paragraph()  # Spacer
    
    doc.add_page_break()

# --- Supplemental Standalone Evidence ---
supp_h = doc.add_paragraph()
supp_h_run = supp_h.add_run("Supplemental Standalone Evidence")
supp_h_run.font.size = Pt(12)
supp_h_run.font.bold = True
supp_h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

supp_p = doc.add_paragraph()
supp_run = supp_p.add_run(
    "This evidence package provides comprehensive technical documentation for all verified Not a Finding items. "
    "Per PIEE PMO STIG Checklist Completion Guide V2.0 Section 4, evidence may be submitted as a comprehensive document "
    "when screenshots are combined or replaced by equivalent technical documentation and compliance statements."
)
supp_run.font.size = Pt(10)
supp_p.paragraph_format.line_spacing = 1.15
supp_p.paragraph_format.space_after = Pt(12)

# --- 5. CUI Markings ---
cui_h = doc.add_paragraph()
cui_h_run = cui_h.add_run("5. CUI Markings")
cui_h_run.font.size = Pt(12)
cui_h_run.font.bold = True
cui_h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

cui_text = doc.add_paragraph()
cui_text_run = cui_text.add_run(
    "The SME POC is responsible for ensuring all submissions are properly marked before submitting to the PIEE PMO. "
    "This may include but is not limited to:\n\n"
    "UNCLASSIFIED//CUI"
)
cui_text_run.font.size = Pt(10)

# Footer
last = doc.add_paragraph()
last_run = last.add_run("UNCLASSIFIED//CUI")
last_run.font.size = Pt(8)
last_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
last.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_dir = r"C:\Users\compj\.openclaw\workspace\output"
docx_path = os.path.join(output_dir, "Appian_ASD_STIG_V6R4_All_NAF_Evidence_Package.docx")
doc.save(docx_path)
print(f"DOCX saved: {docx_path}")

# Convert to PDF
import win32com.client as win32
word = win32.Dispatch('Word.Application')
word.Visible = False
pdf_path = docx_path.replace('.docx', '.pdf')
doc_w = word.Documents.Open(docx_path)
doc_w.SaveAs(pdf_path, FileFormat=17)
doc_w.Close()
word.Quit()
print(f"PDF saved: {pdf_path}")
