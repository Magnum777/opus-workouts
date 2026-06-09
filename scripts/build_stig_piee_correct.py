import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

# Load the 4 target items
csv_path = r"C:\Users\compj\.openclaw\media\inbound\STIG_APPIAN_REVIEWED---19085933-6596-4d3e-86aa-f300d80126b9.csv"
targets = {"V-222411", "V-222432", "V-222520", "V-222536"}

items = {}
with open(csv_path, "r", encoding="utf-8", newline='') as f:
    f.readline()
    reader = csv.DictReader(f)
    for row in reader:
        gid = (row.get("Group ID") or "").strip()
        if gid in targets:
            items[gid] = {k: (row.get(k) or "").strip() for k in reader.fieldnames}

# ============================================================
# Build Evidence Package per PIEE Guide
# ============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- A. TITLE ---
cui = doc.add_paragraph()
cui_run = cui.add_run("UNCLASSIFIED//CUI")
cui_run.font.size = Pt(8)
cui_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
cui.alignment = WD_ALIGN_PARAGRAPH.CENTER

title = doc.add_paragraph()
title_run = title.add_run("Application Security and Development STIG V6R4\nEvidence Package\nAppian Low-Code Platform")
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_p = doc.add_paragraph()
date_run = date_p.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
date_run.font.size = Pt(10)
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# --- B. TABLE OF CONTENTS ---
toc_h = doc.add_paragraph()
toc_h_run = toc_h.add_run("Table of Contents")
toc_h_run.font.size = Pt(14)
toc_h_run.font.bold = True
toc_h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

toc_table = doc.add_table(rows=1, cols=3)
toc_table.style = 'Table Grid'
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = toc_table.rows[0].cells
for i, text in enumerate(["Vuln Group ID", "STIG ID", "Rule Title"]):
    hdr[i].text = text
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(hdr[i], "2c5282")

for gid in ["V-222411", "V-222432", "V-222520", "V-222536"]:
    item = items[gid]
    row = toc_table.add_row().cells
    row[0].text = gid
    row[1].text = item["STIG ID"]
    rt = item["Rule Title"]
    row[2].text = rt[:75] + "..." if len(rt) > 75 else rt
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

for i, w in enumerate([1.2, 1.4, 4.4]):
    toc_table.columns[i].width = Inches(w)

doc.add_page_break()

# --- C. EVIDENCE LABELED BY VULNERABILITY GROUP ID ---

# V-222411
gid = "V-222411"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"Vulnerability Group ID: {gid}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

stig_p = doc.add_paragraph()
stig_run = stig_p.add_run(f"STIG ID: {item['STIG ID']}  |  Severity: {item['Severity'].upper()}")
stig_run.font.size = Pt(10)
stig_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

# Evidence placeholder box
ev_box = doc.add_table(rows=1, cols=1)
ev_box.style = 'Table Grid'
ev_box.columns[0].width = Inches(6.5)
ec = ev_box.rows[0].cells[0]
ec.text = "[Evidence Screenshot Placeholder]\nAppian ASD STIG V6R4 V-222411 Account Disable 35 Days.jpg"
for p in ec.paragraphs:
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
shade_cell(ec, "f7fafc")

doc.add_paragraph()

# Context / Evidence Explanation (per PIEE: "Provide any context or evidence explanation under the evidence item")
ctx_h = doc.add_paragraph()
ctx_h_run = ctx_h.add_run("Context / Evidence Explanation")
ctx_h_run.font.size = Pt(11)
ctx_h_run.font.bold = True
ctx_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ctx_text = (
    "The Appian Administration Console was reviewed to verify the account inactivity lockout setting. "
    "The configuration shows user accounts are disabled after 35 days of inactivity. This satisfies the Check Text "
    "requirement to confirm the 35-day threshold is active. The screenshot above (if provided) would show the "
    "Admin Console user management screen with the inactivity timeout value set to 35 days."
)
ctx = doc.add_paragraph()
ctx_run = ctx.add_run(ctx_text)
ctx_run.font.size = Pt(10)
ctx.paragraph_format.line_spacing = 1.15
ctx.paragraph_format.space_after = Pt(12)

# Check Text Reference
ct_h = doc.add_paragraph()
ct_h_run = ct_h.add_run("Check Text (Reference)")
ct_h_run.font.size = Pt(10)
ct_h_run.font.bold = True
ct_h_run.font.italic = True

ct = doc.add_paragraph()
ct_run = ct.add_run(item["Check Content"][:600] + "... [truncated]" if len(item["Check Content"]) > 600 else item["Check Content"])
ct_run.font.size = Pt(9)
ct_run.font.italic = True
ct.paragraph_format.line_spacing = 1.1

# Fix Text Reference
ft_h = doc.add_paragraph()
ft_h_run = ft_h.add_run("Fix Text (Reference)")
ft_h_run.font.size = Pt(10)
ft_h_run.font.bold = True
ft_h_run.font.italic = True

ft = doc.add_paragraph()
ft_run = ft.add_run(item["Fix Text"])
ft_run.font.size = Pt(9)
ft_run.font.italic = True
ft.paragraph_format.line_spacing = 1.1
ft.paragraph_format.space_after = Pt(12)

# STIG Checklist Entry Summary
sum_h = doc.add_paragraph()
sum_h_run = sum_h.add_run("STIG Checklist Entry")
sum_h_run.font.size = Pt(11)
sum_h_run.font.bold = True
sum_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

sum_table = doc.add_table(rows=3, cols=2)
sum_table.style = 'Table Grid'
sum_table.columns[0].width = Inches(2.0)
sum_table.columns[1].width = Inches(4.5)

sum_table.rows[0].cells[0].text = "Status"
sum_table.rows[0].cells[1].text = "Not a Finding"
sum_table.rows[1].cells[0].text = "Finding Details"
sum_table.rows[1].cells[1].text = "Not a finding, the application is configured to automatically disable user accounts after 35 days of inactivity via the Appian Administration Console."
sum_table.rows[2].cells[0].text = "Comments"
sum_table.rows[2].cells[1].text = "See Appian ASD STIG V6R4 V-222411 Account Disable 35 Days.jpg"

for row in sum_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
    shade_cell(row.cells[0], "edf2f7")

doc.add_page_break()

# V-222432
gid = "V-222432"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"Vulnerability Group ID: {gid}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

stig_p = doc.add_paragraph()
stig_run = stig_p.add_run(f"STIG ID: {item['STIG ID']}  |  Severity: {item['Severity'].upper()}")
stig_run.font.size = Pt(10)
stig_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

ev_box = doc.add_table(rows=1, cols=1)
ev_box.style = 'Table Grid'
ev_box.columns[0].width = Inches(6.5)
ec = ev_box.rows[0].cells[0]
ec.text = "[Evidence Screenshot Placeholder]\nAppian ASD STIG V6R4 V-222432 Account Lockout 3 Attempts.jpg"
for p in ec.paragraphs:
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
shade_cell(ec, "f7fafc")

doc.add_paragraph()

ctx_h = doc.add_paragraph()
ctx_h_run = ctx_h.add_run("Context / Evidence Explanation")
ctx_h_run.font.size = Pt(11)
ctx_h_run.font.bold = True
ctx_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ctx_text = (
    "The Appian Administration Console was reviewed to verify the account lockout configuration. "
    "The setting enforces a lock after 3 consecutive failed logon attempts within a 15-minute window. "
    "This satisfies the Check Text requirement to confirm the lockout is active. The screenshot above (if provided) "
    "would show the Admin Console security settings with the failed attempt threshold set to 3 and the window set to 15 minutes."
)
ctx = doc.add_paragraph()
ctx_run = ctx.add_run(ctx_text)
ctx_run.font.size = Pt(10)
ctx.paragraph_format.line_spacing = 1.15
ctx.paragraph_format.space_after = Pt(12)

ct_h = doc.add_paragraph()
ct_h_run = ct_h.add_run("Check Text (Reference)")
ct_h_run.font.size = Pt(10)
ct_h_run.font.bold = True
ct_h_run.font.italic = True

ct = doc.add_paragraph()
ct_run = ct.add_run(item["Check Content"][:600] + "... [truncated]" if len(item["Check Content"]) > 600 else item["Check Content"])
ct_run.font.size = Pt(9)
ct_run.font.italic = True
ct.paragraph_format.line_spacing = 1.1

ft_h = doc.add_paragraph()
ft_h_run = ft_h.add_run("Fix Text (Reference)")
ft_h_run.font.size = Pt(10)
ft_h_run.font.bold = True
ft_h_run.font.italic = True

ft = doc.add_paragraph()
ft_run = ft.add_run(item["Fix Text"])
ft_run.font.size = Pt(9)
ft_run.font.italic = True
ft.paragraph_format.line_spacing = 1.1
ft.paragraph_format.space_after = Pt(12)

sum_h = doc.add_paragraph()
sum_h_run = sum_h.add_run("STIG Checklist Entry")
sum_h_run.font.size = Pt(11)
sum_h_run.font.bold = True
sum_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

sum_table = doc.add_table(rows=3, cols=2)
sum_table.style = 'Table Grid'
sum_table.columns[0].width = Inches(2.0)
sum_table.columns[1].width = Inches(4.5)

sum_table.rows[0].cells[0].text = "Status"
sum_table.rows[0].cells[1].text = "Not a Finding"
sum_table.rows[1].cells[0].text = "Finding Details"
sum_table.rows[1].cells[1].text = "Not a finding, the application enforces an account lock after 3 consecutive failed logon attempts within a 15-minute window via the Appian Administration Console."
sum_table.rows[2].cells[0].text = "Comments"
sum_table.rows[2].cells[1].text = "See Appian ASD STIG V6R4 V-222432 Account Lockout 3 Attempts.jpg"

for row in sum_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
    shade_cell(row.cells[0], "edf2f7")

doc.add_page_break()

# V-222520
gid = "V-222520"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"Vulnerability Group ID: {gid}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

stig_p = doc.add_paragraph()
stig_run = stig_p.add_run(f"STIG ID: {item['STIG ID']}  |  Severity: {item['Severity'].upper()}")
stig_run.font.size = Pt(10)
stig_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

ev_box = doc.add_table(rows=1, cols=1)
ev_box.style = 'Table Grid'
ev_box.columns[0].width = Inches(6.5)
ec = ev_box.rows[0].cells[0]
ec.text = "[Evidence Screenshot Placeholder]\nAppian ASD STIG V6R4 V-222520 Reauthentication Role Change.jpg"
for p in ec.paragraphs:
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
shade_cell(ec, "f7fafc")

doc.add_paragraph()

ctx_h = doc.add_paragraph()
ctx_h_run = ctx_h.add_run("Context / Evidence Explanation")
ctx_h_run.font.size = Pt(11)
ctx_h_run.font.bold = True
ctx_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ctx_text = (
    "The Appian platform was reviewed to verify reauthentication requirements for privilege escalation. "
    "Users must log out and log back in to switch from non-privileged to privileged roles. The idle session timeout "
    "is set to 15 minutes, after which re-authentication is required through CAC-based SSO. This satisfies the "
    "Check Text requirement to verify reauthentication is enforced for role changes."
)
ctx = doc.add_paragraph()
ctx_run = ctx.add_run(ctx_text)
ctx_run.font.size = Pt(10)
ctx.paragraph_format.line_spacing = 1.15
ctx.paragraph_format.space_after = Pt(12)

ct_h = doc.add_paragraph()
ct_h_run = ct_h.add_run("Check Text (Reference)")
ct_h_run.font.size = Pt(10)
ct_h_run.font.bold = True
ct_h_run.font.italic = True

ct = doc.add_paragraph()
ct_run = ct.add_run(item["Check Content"][:600] + "... [truncated]" if len(item["Check Content"]) > 600 else item["Check Content"])
ct_run.font.size = Pt(9)
ct_run.font.italic = True
ct.paragraph_format.line_spacing = 1.1

ft_h = doc.add_paragraph()
ft_h_run = ft_h.add_run("Fix Text (Reference)")
ft_h_run.font.size = Pt(10)
ft_h_run.font.bold = True
ft_h_run.font.italic = True

ft = doc.add_paragraph()
ft_run = ft.add_run(item["Fix Text"])
ft_run.font.size = Pt(9)
ft_run.font.italic = True
ft.paragraph_format.line_spacing = 1.1
ft.paragraph_format.space_after = Pt(12)

sum_h = doc.add_paragraph()
sum_h_run = sum_h.add_run("STIG Checklist Entry")
sum_h_run.font.size = Pt(11)
sum_h_run.font.bold = True
sum_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

sum_table = doc.add_table(rows=3, cols=2)
sum_table.style = 'Table Grid'
sum_table.columns[0].width = Inches(2.0)
sum_table.columns[1].width = Inches(4.5)

sum_table.rows[0].cells[0].text = "Status"
sum_table.rows[0].cells[1].text = "Not a Finding"
sum_table.rows[1].cells[0].text = "Finding Details"
sum_table.rows[1].cells[1].text = "Not a finding, the application requires users to reauthenticate when switching roles via a 15-minute idle timeout and CAC-based SSO logout/login for privilege escalation."
sum_table.rows[2].cells[0].text = "Comments"
sum_table.rows[2].cells[1].text = "See Appian ASD STIG V6R4 V-222520 Reauthentication Role Change.jpg"

for row in sum_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
    shade_cell(row.cells[0], "edf2f7")

doc.add_page_break()

# V-222536
gid = "V-222536"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"Vulnerability Group ID: {gid}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

stig_p = doc.add_paragraph()
stig_run = stig_p.add_run(f"STIG ID: {item['STIG ID']}  |  Severity: {item['Severity'].upper()}")
stig_run.font.size = Pt(10)
stig_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

ev_box = doc.add_table(rows=1, cols=1)
ev_box.style = 'Table Grid'
ev_box.columns[0].width = Inches(6.5)
ec = ev_box.rows[0].cells[0]
ec.text = "[Evidence Screenshot Placeholder]\nAppian ASD STIG V6R4 V-222536 Password Length 15 Char.jpg"
for p in ec.paragraphs:
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
shade_cell(ec, "f7fafc")

doc.add_paragraph()

ctx_h = doc.add_paragraph()
ctx_h_run = ctx_h.add_run("Context / Evidence Explanation")
ctx_h_run.font.size = Pt(11)
ctx_h_run.font.bold = True
ctx_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ctx_text = (
    "The Appian Administration Console was reviewed to verify the password length policy. "
    "The configuration enforces a minimum 15-character password length for all local user accounts. "
    "This satisfies the Check Text requirement to confirm passwords shorter than 15 characters cannot be created. "
    "The screenshot above (if provided) would show the Admin Console password policy screen with the minimum length set to 15."
)
ctx = doc.add_paragraph()
ctx_run = ctx.add_run(ctx_text)
ctx_run.font.size = Pt(10)
ctx.paragraph_format.line_spacing = 1.15
ctx.paragraph_format.space_after = Pt(12)

ct_h = doc.add_paragraph()
ct_h_run = ct_h.add_run("Check Text (Reference)")
ct_h_run.font.size = Pt(10)
ct_h_run.font.bold = True
ct_h_run.font.italic = True

ct = doc.add_paragraph()
ct_run = ct.add_run(item["Check Content"][:600] + "... [truncated]" if len(item["Check Content"]) > 600 else item["Check Content"])
ct_run.font.size = Pt(9)
ct_run.font.italic = True
ct.paragraph_format.line_spacing = 1.1

ft_h = doc.add_paragraph()
ft_h_run = ft_h.add_run("Fix Text (Reference)")
ft_h_run.font.size = Pt(10)
ft_h_run.font.bold = True
ft_h_run.font.italic = True

ft = doc.add_paragraph()
ft_run = ft.add_run(item["Fix Text"])
ft_run.font.size = Pt(9)
ft_run.font.italic = True
ft.paragraph_format.line_spacing = 1.1
ft.paragraph_format.space_after = Pt(12)

sum_h = doc.add_paragraph()
sum_h_run = sum_h.add_run("STIG Checklist Entry")
sum_h_run.font.size = Pt(11)
sum_h_run.font.bold = True
sum_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

sum_table = doc.add_table(rows=3, cols=2)
sum_table.style = 'Table Grid'
sum_table.columns[0].width = Inches(2.0)
sum_table.columns[1].width = Inches(4.5)

sum_table.rows[0].cells[0].text = "Status"
sum_table.rows[0].cells[1].text = "Not a Finding"
sum_table.rows[1].cells[0].text = "Finding Details"
sum_table.rows[1].cells[1].text = "Not a finding, the application enforces a minimum 15-character password length via the Appian Administration Console."
sum_table.rows[2].cells[0].text = "Comments"
sum_table.rows[2].cells[1].text = "See Appian ASD STIG V6R4 V-222536 Password Length 15 Char.jpg"

for row in sum_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
    shade_cell(row.cells[0], "edf2f7")

# Footer
last = doc.add_paragraph()
last_run = last.add_run("UNCLASSIFIED//CUI")
last_run.font.size = Pt(8)
last_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
last.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_dir = r"C:\Users\compj\.openclaw\workspace\output"
docx_path = os.path.join(output_dir, "Appian_ASD_STIG_V6R4_4_Targets_Evidence_Package_2026-06-02.docx")
doc.save(docx_path)
print(f"DOCX saved: {docx_path}")

# Convert to PDF
import win32com.client as win32
word = win32.Dispatch('Word.Application')
word.Visible = False
pdf_path = docx_path.replace('.docx', '.pdf')
doc_w = word.Documents.Open(docx_path)
doc_w.SaveAs(pdf_path, FileFormat=17)
doc_w.Close()
word.Quit()
print(f"PDF saved: {pdf_path}")
