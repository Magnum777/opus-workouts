import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_shaded_label(doc, label, value, color_hex="2c5282", font_color="FFFFFF"):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    c0, c1 = table.rows[0].cells
    c0.text = label
    c1.text = value
    shade_cell(c0, color_hex)
    for p in c0.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for p in c1.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)

def add_open_item(doc, item, page_num):
    gid = item['Group ID']
    sid = item['STIG ID']
    sev = (item['Severity'] or '').upper()
    title = item['Rule Title']
    fix_text = item['Fix Text']
    check_text = item['Check Content']
    comments = item['Comments']
    
    # Determine Admin Console setting name
    rt_lower = title.lower()
    comments_lower = comments.lower()
    
    # Map to likely Admin Console section
    admin_section = ""
    if 'password' in rt_lower:
        if 'length' in rt_lower:
            admin_section = "Admin Console > Authentication > Password Format > Minimum password length"
        elif 'complexity' in rt_lower or 'uppercase' in rt_lower or 'lowercase' in rt_lower or 'number' in rt_lower or 'special' in rt_lower:
            admin_section = "Admin Console > Authentication > Password Format > Password complexity requirements"
        elif 'expiration' in rt_lower:
            admin_section = "Admin Console > Authentication > Password Expiration"
        elif 'temporary' in rt_lower:
            admin_section = "Admin Console > Authentication > Temporary Passwords"
        elif 'history' in rt_lower:
            admin_section = "Admin Console > Authentication > Password History"
        elif 'transmit' in rt_lower or 'transmission' in rt_lower or 'cryptographically' in rt_lower:
            admin_section = "Enable TLS 1.2+ in Appian Cloud / reverse proxy configuration"
        else:
            admin_section = "Admin Console > Authentication > Password settings"
    elif 'session' in rt_lower or 'concurrent' in rt_lower or 'logon session' in rt_lower:
        admin_section = "Admin Console > Authentication > Session Management / Concurrent Sessions"
    elif 'idle' in rt_lower or 'timeout' in rt_lower or 'terminate' in rt_lower:
        admin_section = "Admin Console > Authentication > Session Timeout"
    elif 'account lock' in rt_lower or 'lockout' in rt_lower:
        admin_section = "Admin Console > Authentication > Account Locking"
    elif 'deactivat' in rt_lower or 'disable' in rt_lower or 'inactive' in rt_lower:
        admin_section = "Admin Console > Authentication > Account Deactivation"
    elif 'temporary account' in rt_lower:
        admin_section = "Admin Console > User Management > Temporary Account Provisioning"
    elif 'saml' in rt_lower or 'notbefore' in rt_lower or 'notonorafter' in rt_lower or 'onetimeuse' in rt_lower:
        admin_section = "Admin Console > Authentication > SAML > SAML Assertion Configuration"
    elif 'cac' in rt_lower or 'pki' in rt_lower or 'piv' in rt_lower or 'certificate' in rt_lower:
        admin_section = "Admin Console > Authentication > SAML/External Auth > Certificate-based Authentication"
    elif 'multi-factor' in rt_lower or 'multifactor' in rt_lower or 'mfa' in rt_lower:
        admin_section = "Admin Console > Authentication > Multi-factor Authentication"
    elif 'fips' in rt_lower:
        admin_section = "Appian Cloud/Server config: Enable FIPS 140-2 mode (requires support case for Cloud)"
    elif 'branding' in rt_lower or 'banner' in rt_lower or 'mark' in rt_lower or 'classification' in rt_lower:
        admin_section = "Admin Console > Branding > Site Banner / Classification Banner"
    elif 'remember me' in rt_lower or 'remember-me' in rt_lower:
        admin_section = "Admin Console > Authentication > Remember Me (DISABLE per STIG)"
    elif 'mutual' in rt_lower or 'endpoint device' in rt_lower:
        admin_section = "Admin Console > Certificates > Mutual TLS / Client Certificate Authentication"
    elif 'tls' in rt_lower or 'encrypt' in rt_lower or 'ssl' in rt_lower:
        admin_section = "TLS Configuration: Admin Console > Security / Cloud Support Case for TLS 1.2 strict"
    elif 'replay' in rt_lower:
        admin_section = "Admin Console > Authentication > SAML > Replay-resistant nonce/timestamp"
    else:
        admin_section = "Admin Console > Review applicable Authentication/Security settings"
    
    # Vuln ID header (blue)
    h = doc.add_paragraph()
    h_run = h.add_run(gid)
    h_run.font.size = Pt(14)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    h.paragraph_format.space_after = Pt(4)
    
    # STIG ID and Severity
    stig_p = doc.add_paragraph()
    stig_run = stig_p.add_run(f"STIG ID: {sid}  |  Severity: {sev}")
    stig_run.font.size = Pt(10)
    stig_run.font.bold = True
    
    # Status: OPEN (red)
    status_p = doc.add_paragraph()
    status_label = status_p.add_run("Status: ")
    status_label.font.size = Pt(10)
    status_label.font.bold = True
    status_val = status_p.add_run("OPEN — FIX VIA ADMIN CONSOLE")
    status_val.font.size = Pt(10)
    status_val.font.bold = True
    status_val.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    status_p.paragraph_format.space_after = Pt(4)
    
    # Rule Title
    rt = doc.add_paragraph()
    rt_run = rt.add_run(f"Rule: {title}")
    rt_run.font.size = Pt(10)
    rt.paragraph_format.space_after = Pt(6)
    
    # Admin Console Fix Location
    fix_loc = doc.add_paragraph()
    fix_loc_label = fix_loc.add_run("Admin Console Fix Location: ")
    fix_loc_label.font.size = Pt(10)
    fix_loc_label.font.bold = True
    fix_loc_val = fix_loc.add_run(admin_section)
    fix_loc_val.font.size = Pt(10)
    fix_loc_val.font.italic = True
    fix_loc.paragraph_format.space_after = Pt(6)
    
    # Evidence placeholder (screenshot after fix)
    ev_box = doc.add_table(rows=1, cols=1)
    ev_box.style = 'Table Grid'
    ev_box.columns[0].width = Inches(6.5)
    ec = ev_box.rows[0].cells[0]
    
    # Determine short descriptor
    short_desc = admin_section.split('>')[-1].strip()[:40]
    ev_filename = f"Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    ec.text = f"[FIXED — Screenshot Placeholder]\n\n{ev_filename}\n\n(Take screenshot after fixing in Admin Console, then replace this placeholder)"
    for p in ec.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    shade_cell(ec, "fff5f5")
    
    doc.add_paragraph()
    
    # Fix Text reference
    fix_h = doc.add_paragraph()
    fix_h_run = fix_h.add_run("Fix Text (from STIG):")
    fix_h_run.font.size = Pt(10)
    fix_h_run.font.bold = True
    fix_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
    
    fix_p = doc.add_paragraph()
    fix_trunc = fix_text[:500]
    if len(fix_text) > 500:
        fix_trunc += "..."
    fix_run = fix_p.add_run(fix_trunc)
    fix_run.font.size = Pt(9)
    fix_p.paragraph_format.line_spacing = 1.15
    fix_p.paragraph_format.space_after = Pt(6)
    
    # Check Text reference
    check_h = doc.add_paragraph()
    check_h_run = check_h.add_run("Check Text (from STIG):")
    check_h_run.font.size = Pt(10)
    check_h_run.font.bold = True
    check_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
    
    check_p = doc.add_paragraph()
    check_trunc = check_text[:500]
    if len(check_text) > 500:
        check_trunc += "..."
    check_run = check_p.add_run(check_trunc)
    check_run.font.size = Pt(9)
    check_p.paragraph_format.line_spacing = 1.15
    check_p.paragraph_format.space_after = Pt(6)
    
    # Current Comment (from CSV)
    if comments.strip():
        comm_h = doc.add_paragraph()
        comm_h_run = comm_h.add_run("Current CSV Comment:")
        comm_h_run.font.size = Pt(10)
        comm_h_run.font.bold = True
        comm_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
        
        comm_p = doc.add_paragraph()
        comm_run = comm_p.add_run(comments[:400])
        comm_run.font.size = Pt(9)
        comm_run.font.italic = True
        comm_p.paragraph_format.line_spacing = 1.15
        comm_p.paragraph_format.space_after = Pt(6)
    
    # After-fix checklist entry (what to put in .cklb once fixed)
    cklb_h = doc.add_paragraph()
    cklb_h_run = cklb_h.add_run("STIG Checklist Entry (after fix):")
    cklb_h_run.font.size = Pt(11)
    cklb_h_run.font.bold = True
    cklb_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
    
    cklb_table = doc.add_table(rows=3, cols=2)
    cklb_table.style = 'Table Grid'
    cklb_table.columns[0].width = Inches(2.0)
    cklb_table.columns[1].width = Inches(4.5)
    
    cklb_table.rows[0].cells[0].text = "Status"
    cklb_table.rows[0].cells[1].text = "Not a Finding"
    cklb_table.rows[1].cells[0].text = "Finding Details"
    cklb_table.rows[1].cells[1].text = f"Not a finding, the application is configured to meet the requirement per the {admin_section}. See evidence screenshot for configuration verification."
    cklb_table.rows[2].cells[0].text = "Comments"
    cklb_table.rows[2].cells[1].text = f"See Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    
    for row in cklb_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        shade_cell(row.cells[0], "edf2f7")
    
    doc.add_page_break()
    return page_num + 1

# Load ALL items from CSV
csv_path = r"C:\Users\compj\.openclaw\media\inbound\STIG_APPIAN_REVIEWED---c4f9d089-0337-4c23-8afd-fd0f631e28a0.csv"
all_items = {}
open_items = []
notafinding_items = []
with open(csv_path, "r", encoding="utf-8", newline='') as f:
    f.readline()
    reader = csv.DictReader(f)
    for row in reader:
        gid = (row.get("Group ID") or "").strip()
        status = (row.get("Status") or "").strip().lower()
        all_items[gid] = row
        if status == 'open':
            open_items.append(row)
        elif status == 'not a finding':
            notafinding_items.append(row)

print(f"Total items: {len(all_items)}, Open: {len(open_items)}, Not a Finding: {len(notafinding_items)}")

# Identify Admin Console fixable
admin_console_keywords = [
    'password', 'session', 'timeout', 'concurrent', 'logon session', 'idle', 'account lock',
    'deactivation', 'deactivated', 'inactive', 'temporary password', 'temporary account',
    'remember me', 'saml', 'multi-factor', 'mfa', 'branding', 'classification', 'banner',
    'site banner', 'admin console', 'authentication', 'login', 'logon', 'lockout',
    'fips', 'tls', 'encrypt', 'mutual', 'endpoint device', 'certificate', 'cac', 'pki', 'piv'
]

admin_items = []
other_open_items = []
for item in open_items:
    rt = (item.get('Rule Title') or '').lower()
    comments = (item.get('Comments') or '').lower()
    fix = (item.get('Fix Text') or '').lower()
    combined = rt + ' ' + comments + ' ' + fix
    
    if any(k in combined for k in admin_console_keywords):
        admin_items.append(item)
    else:
        other_open_items.append(item)

print(f"Admin Console fixable: {len(admin_items)}, Other Open: {len(other_open_items)}")

# ============================================================
# Build full evidence package
# ============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- A. TITLE ---
title = doc.add_paragraph()
title_run = title.add_run("Application Security and Development STIG\nVersion: 6, Release: 4")
title_run.font.size = Pt(20)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

cui = doc.add_paragraph()
cui_run = cui.add_run("UNCLASSIFIED//CUI")
cui_run.font.size = Pt(10)
cui_run.font.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("Appian Low-Code Platform — Full Evidence Package\nIncludes: 4 Not a Finding (verified) + 30 Open (fixable via Admin Console)")
subtitle_run.font.size = Pt(11)
subtitle_run.font.italic = True

date_p = doc.add_paragraph()
date_run = date_p.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
date_run.font.size = Pt(10)

doc.add_paragraph()

# --- B. CONTENTS ---
contents_h = doc.add_paragraph()
contents_h_run = contents_h.add_run("Contents")
contents_h_run.font.size = Pt(14)
contents_h_run.font.bold = True
contents_h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

toc_items = []
page_num = 3

# Not a Finding items first
for gid in ["V-222411", "V-222432", "V-222520", "V-222536"]:
    if gid in all_items:
        toc_items.append((gid, page_num, "Not a Finding"))
        page_num += 1

# Admin Console fixable
for item in admin_items:
    toc_items.append((item['Group ID'], page_num, "OPEN — Admin Console Fix"))
    page_num += 1

for gid, pg, status in toc_items:
    toc_line = doc.add_paragraph()
    toc_run = toc_line.add_run(f"{gid} ({status})")
    toc_run.font.size = Pt(11)
    dots = " ." * (70 - len(gid) - len(status) - 5)
    dots_run = toc_line.add_run(f"{dots} {pg}")
    dots_run.font.size = Pt(11)
    toc_line.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# NOT A FINDING ITEMS (verified, evidence + checklist ref)
# ============================================================
naf_header = doc.add_paragraph()
naf_header_run = naf_header.add_run("SECTION 1 — NOT A FINDING (Verified)")
naf_header_run.font.size = Pt(16)
naf_header_run.font.bold = True
naf_header_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
doc.add_paragraph()

# The 4 target items
for gid in ["V-222411", "V-222432", "V-222520", "V-222536"]:
    item = all_items[gid]
    
    h = doc.add_paragraph()
    h_run = h.add_run(f"Vulnerability Group ID: {gid}")
    h_run.font.size = Pt(14)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    h.paragraph_format.space_after = Pt(4)
    
    add_shaded_label(doc, "Status", "Not a Finding", "008000")
    add_shaded_label(doc, "STIG ID", item['STIG ID'], "2c5282")
    add_shaded_label(doc, "Severity", item['Severity'].upper(), "2c5282")
    
    rt = doc.add_paragraph()
    rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
    rt_run.font.size = Pt(10)
    rt.paragraph_format.space_after = Pt(6)
    
    # Evidence placeholder
    short_desc_map = {
        "V-222411": "Account Disable 35 Days",
        "V-222432": "Account Lockout 3 Attempts",
        "V-222520": "Reauthentication Role Change",
        "V-222536": "Password Length 15 Char"
    }
    short_desc = short_desc_map[gid]
    
    ev_box = doc.add_table(rows=1, cols=1)
    ev_box.style = 'Table Grid'
    ev_box.columns[0].width = Inches(6.5)
    ec = ev_box.rows[0].cells[0]
    ev_filename = f"Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    ec.text = f"[Evidence Screenshot Placeholder]\n\n{ev_filename}"
    for p in ec.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    shade_cell(ec, "f5f5f5")
    
    doc.add_paragraph()
    
    # Context/Explanation
    ctx_h = doc.add_paragraph()
    ctx_h_run = ctx_h.add_run("Context / Evidence Explanation")
    ctx_h_run.font.size = Pt(10)
    ctx_h_run.font.bold = True
    ctx_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
    
    explanations = {
        "V-222411": "The Appian Administration Console user management screen was reviewed to verify the account inactivity lockout setting. The configuration shows user accounts are disabled after 35 days of inactivity.",
        "V-222432": "The Appian Administration Console security settings screen was reviewed to verify the account lockout configuration. The setting enforces a lock after 3 consecutive failed logon attempts within a 15-minute window.",
        "V-222520": "The Appian platform session timeout and role change workflow was reviewed. Users must log out and log back in to switch roles. The idle session timeout is set to 15 minutes with CAC-based SSO.",
        "V-222536": "The Appian Administration Console password policy screen was reviewed. The configuration enforces a minimum 15-character password length for all local user accounts."
    }
    ctx = doc.add_paragraph()
    ctx_run = ctx.add_run(explanations[gid])
    ctx_run.font.size = Pt(10)
    ctx.paragraph_format.line_spacing = 1.15
    ctx.paragraph_format.space_after = Pt(6)
    
    # Checklist reference
    cklb_h = doc.add_paragraph()
    cklb_h_run = cklb_h.add_run("STIG Checklist Entry Reference")
    cklb_h_run.font.size = Pt(10)
    cklb_h_run.font.bold = True
    cklb_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
    
    cklb_table = doc.add_table(rows=3, cols=2)
    cklb_table.style = 'Table Grid'
    cklb_table.columns[0].width = Inches(2.0)
    cklb_table.columns[1].width = Inches(4.5)
    
    cklb_table.rows[0].cells[0].text = "Status"
    cklb_table.rows[0].cells[1].text = "Not a Finding"
    cklb_table.rows[1].cells[0].text = "Finding Details"
    cklb_table.rows[1].cells[1].text = item.get('Finding Details', 'Not a finding, the application is configured to meet the requirement.')[:300]
    cklb_table.rows[2].cells[0].text = "Comments"
    cklb_table.rows[2].cells[1].text = f"See Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    
    for row in cklb_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        shade_cell(row.cells[0], "edf2f7")
    
    doc.add_page_break()

# ============================================================
# OPEN ITEMS — ADMIN CONSOLE FIXABLE
# ============================================================
open_header = doc.add_paragraph()
open_header_run = open_header.add_run("SECTION 2 — OPEN ITEMS FIXABLE VIA ADMIN CONSOLE")
open_header_run.font.size = Pt(16)
open_header_run.font.bold = True
open_header_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph()

intro = doc.add_paragraph()
intro_run = intro.add_run(
    "The following items are currently OPEN (non-compliant) and can be remediated through the "
    "Appian Administration Console. For each item, navigate to the specified Admin Console "
    "section, apply the required configuration, then capture a screenshot of the fixed setting "
    "and replace the placeholder below. Use the 'STIG Checklist Entry (after fix)' table to "
    "populate the .cklb file once verified."
)
intro_run.font.size = Pt(10)
intro.paragraph_format.line_spacing = 1.15
intro.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

page_num = 3 + 4  # After TOC + 4 NAF pages
for item in admin_items:
    page_num = add_open_item(doc, item, page_num)

# Footer
last = doc.add_paragraph()
last_run = last.add_run("UNCLASSIFIED//CUI")
last_run.font.size = Pt(8)
last_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
last.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_dir = r"C:\Users\compj\.openclaw\workspace\output"
docx_path = os.path.join(output_dir, "Appian_ASD_STIG_V6R4_Full_Evidence_Package.docx")
doc.save(docx_path)
print(f"DOCX saved: {docx_path}")

# Convert to PDF
import win32com.client as win32
word = win32.Dispatch('Word.Application')
word.Visible = False
pdf_path = docx_path.replace('.docx', '.pdf')
doc_w = word.Documents.Open(docx_path)
doc_w.SaveAs(pdf_path, FileFormat=17)
doc_w.Close()
word.Quit()
print(f"PDF saved: {pdf_path}")
