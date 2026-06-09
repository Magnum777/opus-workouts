import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_shaded_label(doc, label, value, color_hex="2c5282"):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.5)
    c0, c1 = table.rows[0].cells
    c0.text = label
    c1.text = value
    shade_cell(c0, color_hex)
    for p in c0.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for p in c1.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)

# Load ALL items from CSV
csv_path = r"C:\Users\compj\.openclaw\media\inbound\STIG_APPIAN_REVIEWED---c4f9d089-0337-4c23-8afd-fd0f631e28a0.csv"
all_items = {}
open_items = []
with open(csv_path, "r", encoding="utf-8", newline='') as f:
    f.readline()
    reader = csv.DictReader(f)
    for row in reader:
        gid = (row.get("Group ID") or "").strip()
        status = (row.get("Status") or "").strip().lower()
        all_items[gid] = row
        if status == 'open':
            open_items.append(row)

# STRICT filter: only true toggle/value changes in Admin Console
# YES: session timeout, concurrent sessions, password length/complexity, account lockout,
#      idle timeout, account deactivation days, remember me toggle, temporary password,
#      password history, branding/classification banner text
# NO: SAML, CAC/PKI, TLS, FIPS, mutual SSL, SIEM streaming, cert configs, anything requiring external setup
toggle_value_keywords = [
    'session timeout', 'concurrent session', 'concurrent logon', 'idle time', 'idle timeout',
    'password length', 'minimum password length', 'password complexity', 'password format',
    'account lock', 'lockout', 'lock out', 'failed logon',
    'deactivat', 'disable account', 'inactive account', 'inactivity',
    'remember me', 'remember-me',
    'temporary password', 'temp password',
    'password history', 'password reuse',
    'branding', 'classification banner', 'site banner', 'banner text',
    'number of logon sessions', 'limit the number of logon'
]
# Exclude anything that requires external systems
toggle_value_items = []
for item in open_items:
    rt = (item.get('Rule Title') or '').lower()
    comments = (item.get('Comments') or '').lower()
    fix = (item.get('Fix Text') or '').lower()
    combined = rt + ' ' + comments + ' ' + fix
    
    # Must match a toggle/value keyword
    is_toggle_value = any(k in combined for k in toggle_value_keywords)
    
    # Must NOT require external setup (SAML, CAC, TLS, FIPS, SIEM, cert, mutual, etc.)
    external_keywords = ['saml', 'cac', 'pki', 'piv', 'tls', 'fips', 'certificate', 'mutual',
                         'siem', 'syslog', 'external authentication', 'idp', 'identity provider',
                         'soap', 'ws-security', 'cryptographic module', 'penetration test',
                         'contingency', 'backup', 'recovery plan', 'disaster recovery']
    requires_external = any(k in combined for k in external_keywords)
    
    if is_toggle_value and not requires_external:
        toggle_value_items.append(item)

print(f"Total Open: {len(open_items)}")
print(f"Toggle/Value fixable in Admin Console: {len(toggle_value_items)}")
for item in toggle_value_items:
    print(f"  {item['Group ID']} | {item['STIG ID']} | {item['Rule Title'][:70]}...")

# ============================================================
# Build lean evidence package
# ============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- A. TITLE ---
title = doc.add_paragraph()
title_run = title.add_run("Application Security and Development STIG\nVersion: 6, Release: 4")
title_run.font.size = Pt(20)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

cui = doc.add_paragraph()
cui_run = cui.add_run("UNCLASSIFIED//CUI")
cui_run.font.size = Pt(10)
cui_run.font.bold = True

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run("Appian Low-Code Platform — Admin Console Fixable Evidence Package\nNot a Finding (verified) + Open (toggle/value fixes)")
subtitle_run.font.size = Pt(11)
subtitle_run.font.italic = True

date_p = doc.add_paragraph()
date_run = date_p.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
date_run.font.size = Pt(10)

doc.add_paragraph()

# --- B. CONTENTS ---
contents_h = doc.add_paragraph()
contents_h_run = contents_h.add_run("Contents")
contents_h_run.font.size = Pt(14)
contents_h_run.font.bold = True
contents_h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

toc_items = []
page_num = 3

# Not a Finding items (verified)
for gid in ["V-222411", "V-222432", "V-222520", "V-222536"]:
    toc_items.append((gid, page_num, "Not a Finding"))
    page_num += 1

# Toggle/Value fixable items
for item in toggle_value_items:
    toc_items.append((item['Group ID'], page_num, "OPEN — Toggle/Value Fix"))
    page_num += 1

for gid, pg, status in toc_items:
    toc_line = doc.add_paragraph()
    toc_run = toc_line.add_run(f"{gid} ({status})")
    toc_run.font.size = Pt(11)
    dots = " ." * (70 - len(gid) - len(status) - 5)
    dots_run = toc_line.add_run(f"{dots} {pg}")
    dots_run.font.size = Pt(11)
    toc_line.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# SECTION 1 — NOT A FINDING (Verified)
# ============================================================
naf_header = doc.add_paragraph()
naf_header_run = naf_header.add_run("SECTION 1 — NOT A FINDING (Verified)")
naf_header_run.font.size = Pt(16)
naf_header_run.font.bold = True
naf_header_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
doc.add_paragraph()

naf_intro = doc.add_paragraph()
naf_intro_run = naf_intro.add_run(
    "The following items have been verified as compliant. Evidence screenshots are provided below. "
    "Copy the Finding Details and Comments into the STIG checklist .cklb file."
)
naf_intro_run.font.size = Pt(10)
naf_intro.paragraph_format.line_spacing = 1.15
naf_intro.paragraph_format.space_after = Pt(12)

short_desc_map = {
    "V-222411": "Account Disable 35 Days",
    "V-222432": "Account Lockout 3 Attempts",
    "V-222520": "Reauthentication Role Change",
    "V-222536": "Password Length 15 Char"
}

explanations = {
    "V-222411": "The Appian Administration Console user management screen was reviewed to verify the account inactivity lockout setting. The configuration shows user accounts are disabled after 35 days of inactivity.",
    "V-222432": "The Appian Administration Console security settings screen was reviewed to verify the account lockout configuration. The setting enforces a lock after 3 consecutive failed logon attempts within a 15-minute window.",
    "V-222520": "The Appian platform session timeout and role change workflow was reviewed. Users must log out and log back in to switch roles. The idle session timeout is set to 15 minutes with CAC-based SSO.",
    "V-222536": "The Appian Administration Console password policy screen was reviewed. The configuration enforces a minimum 15-character password length for all local user accounts."
}

for gid in ["V-222411", "V-222432", "V-222520", "V-222536"]:
    item = all_items[gid]
    short_desc = short_desc_map[gid]
    
    # Vuln ID header
    h = doc.add_paragraph()
    h_run = h.add_run(f"Vulnerability Group ID: {gid}")
    h_run.font.size = Pt(14)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    h.paragraph_format.space_after = Pt(4)
    
    add_shaded_label(doc, "Status", "Not a Finding", "008000")
    add_shaded_label(doc, "STIG ID", item['STIG ID'])
    add_shaded_label(doc, "Severity", item['Severity'].upper())
    
    rt = doc.add_paragraph()
    rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
    rt_run.font.size = Pt(10)
    rt.paragraph_format.space_after = Pt(6)
    
    # Evidence placeholder
    ev_box = doc.add_table(rows=1, cols=1)
    ev_box.style = 'Table Grid'
    ev_box.columns[0].width = Inches(6.5)
    ec = ev_box.rows[0].cells[0]
    ev_filename = f"Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    ec.text = f"[Evidence Screenshot Placeholder]\n\n{ev_filename}"
    for p in ec.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    shade_cell(ec, "f5f5f5")
    
    doc.add_paragraph()
    
    # Explanation
    exp_p = doc.add_paragraph()
    exp_label = exp_p.add_run("Explanation/Context: ")
    exp_label.font.size = Pt(10)
    exp_label.font.bold = True
    exp_text = exp_p.add_run(f"({explanations[gid]})")
    exp_text.font.size = Pt(10)
    exp_p.paragraph_format.line_spacing = 1.2
    exp_p.paragraph_format.space_after = Pt(6)
    
    # Checklist ref
    cklb_h = doc.add_paragraph()
    cklb_h_run = cklb_h.add_run("STIG Checklist Entry Reference")
    cklb_h_run.font.size = Pt(10)
    cklb_h_run.font.bold = True
    cklb_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
    
    cklb_table = doc.add_table(rows=3, cols=2)
    cklb_table.style = 'Table Grid'
    cklb_table.columns[0].width = Inches(2.0)
    cklb_table.columns[1].width = Inches(4.5)
    cklb_table.rows[0].cells[0].text = "Status"
    cklb_table.rows[0].cells[1].text = "Not a Finding"
    cklb_table.rows[1].cells[0].text = "Finding Details"
    cklb_table.rows[1].cells[1].text = item.get('Finding Details', f"Not a finding, the application is configured to meet the requirement per {short_desc}.")[:350]
    cklb_table.rows[2].cells[0].text = "Comments"
    cklb_table.rows[2].cells[1].text = f"See Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    for row in cklb_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        shade_cell(row.cells[0], "edf2f7")
    
    doc.add_page_break()

# ============================================================
# SECTION 2 — OPEN (Toggle/Value Fixes in Admin Console)
# ============================================================
open_header = doc.add_paragraph()
open_header_run = open_header.add_run("SECTION 2 — OPEN (Admin Console Toggle/Value Fixes)")
open_header_run.font.size = Pt(16)
open_header_run.font.bold = True
open_header_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
doc.add_paragraph()

open_intro = doc.add_paragraph()
open_intro_run = open_intro.add_run(
    "The following items are OPEN and can be remediated by changing a toggle, checkbox, or numeric value "
    "in the Appian Administration Console. No external systems, certificates, or IdP configuration is required. "
    "Navigate to the specified Admin Console section, apply the required setting, capture a screenshot, "
    "and replace the placeholder below. Then copy the 'STIG Checklist Entry (after fix)' into the .cklb file."
)
open_intro_run.font.size = Pt(10)
open_intro.paragraph_format.line_spacing = 1.15
open_intro.paragraph_format.space_after = Pt(12)

doc.add_paragraph()

# Map each toggle/value item to its Admin Console location and short descriptor
for item in toggle_value_items:
    gid = item['Group ID']
    sid = item['STIG ID']
    sev = (item['Severity'] or '').upper()
    title = item['Rule Title']
    fix_text = item['Fix Text']
    check_text = item['Check Content']
    comments = item['Comments']
    
    rt_lower = title.lower()
    comments_lower = comments.lower()
    
    # Determine exact Admin Console section
    if 'concurrent' in rt_lower or 'number of logon session' in rt_lower:
        admin_section = "Admin Console > Authentication > Session Management > Max Concurrent Sessions"
        short_desc = "Concurrent Sessions"
    elif 'idle' in rt_lower and ('admin' in rt_lower or 'privileged' in rt_lower):
        admin_section = "Admin Console > Authentication > Session Timeout > Admin Idle Timeout"
        short_desc = "Admin Idle Timeout"
    elif 'idle' in rt_lower or 'timeout' in rt_lower:
        admin_section = "Admin Console > Authentication > Session Timeout > User Idle Timeout"
        short_desc = "User Idle Timeout"
    elif 'password length' in rt_lower or 'minimum password length' in rt_lower:
        admin_section = "Admin Console > Authentication > Password Format > Minimum Password Length"
        short_desc = "Min Password Length"
    elif 'password complexity' in rt_lower or 'uppercase' in rt_lower or 'lowercase' in rt_lower or 'number' in rt_lower or 'special' in rt_lower:
        admin_section = "Admin Console > Authentication > Password Format > Password Complexity Requirements"
        short_desc = "Password Complexity"
    elif 'password history' in rt_lower or 'password reuse' in rt_lower:
        admin_section = "Admin Console > Authentication > Password Format > Password History"
        short_desc = "Password History"
    elif 'temporary password' in rt_lower or 'temp password' in rt_lower:
        admin_section = "Admin Console > Authentication > Password Expiration > Temporary Password Expiry"
        short_desc = "Temp Password Expiry"
    elif 'account lock' in rt_lower or 'lockout' in rt_lower or 'lock out' in rt_lower:
        admin_section = "Admin Console > Authentication > Account Locking > Failed Login Attempts"
        short_desc = "Account Lockout"
    elif 'deactivat' in rt_lower or 'disable account' in rt_lower or 'inactive' in rt_lower or 'inactivity' in rt_lower:
        admin_section = "Admin Console > Authentication > Account Deactivation > Days of Inactivity"
        short_desc = "Account Deactivation"
    elif 'remember me' in rt_lower or 'remember-me' in rt_lower:
        admin_section = "Admin Console > Authentication > Remember Me > DISABLE (toggle off)"
        short_desc = "Disable Remember Me"
    elif 'branding' in rt_lower or 'classification banner' in rt_lower or 'site banner' in rt_lower or 'banner text' in rt_lower:
        admin_section = "Admin Console > Branding > Site Banner / Classification Banner Text"
        short_desc = "Classification Banner"
    else:
        admin_section = "Admin Console > Authentication > Review applicable settings"
        short_desc = "Security Setting"
    
    # Vuln ID header
    h = doc.add_paragraph()
    h_run = h.add_run(f"Vulnerability Group ID: {gid}")
    h_run.font.size = Pt(14)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    h.paragraph_format.space_after = Pt(4)
    
    # Status: OPEN
    status_p = doc.add_paragraph()
    status_label = status_p.add_run("Status: ")
    status_label.font.size = Pt(10)
    status_label.font.bold = True
    status_val = status_p.add_run("OPEN — FIX VIA ADMIN CONSOLE")
    status_val.font.size = Pt(10)
    status_val.font.bold = True
    status_val.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    status_p.paragraph_format.space_after = Pt(4)
    
    add_shaded_label(doc, "STIG ID", sid)
    add_shaded_label(doc, "Severity", sev)
    
    rt = doc.add_paragraph()
    rt_run = rt.add_run(f"Rule: {title}")
    rt_run.font.size = Pt(10)
    rt.paragraph_format.space_after = Pt(6)
    
    # Admin Console Fix Location
    fix_loc = doc.add_paragraph()
    fix_loc_label = fix_loc.add_run("Admin Console Fix Location: ")
    fix_loc_label.font.size = Pt(10)
    fix_loc_label.font.bold = True
    fix_loc_val = fix_loc.add_run(admin_section)
    fix_loc_val.font.size = Pt(10)
    fix_loc_val.font.italic = True
    fix_loc.paragraph_format.space_after = Pt(6)
    
    # Evidence placeholder (screenshot AFTER fix)
    ev_box = doc.add_table(rows=1, cols=1)
    ev_box.style = 'Table Grid'
    ev_box.columns[0].width = Inches(6.5)
    ec = ev_box.rows[0].cells[0]
    ev_filename = f"Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    ec.text = f"[FIXED — Screenshot Placeholder]\n\n{ev_filename}\n\n(Take screenshot after fixing in Admin Console, then replace this placeholder)"
    for p in ec.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    shade_cell(ec, "fff5f5")
    
    doc.add_paragraph()
    
    # Fix Text
    fix_h = doc.add_paragraph()
    fix_h_run = fix_h.add_run("Fix Text (from STIG):")
    fix_h_run.font.size = Pt(10)
    fix_h_run.font.bold = True
    fix_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
    
    fix_p = doc.add_paragraph()
    fix_trunc = fix_text[:500]
    if len(fix_text) > 500:
        fix_trunc += "..."
    fix_run = fix_p.add_run(fix_trunc)
    fix_run.font.size = Pt(9)
    fix_p.paragraph_format.line_spacing = 1.15
    fix_p.paragraph_format.space_after = Pt(6)
    
    # Current CSV Comment
    if comments.strip():
        comm_h = doc.add_paragraph()
        comm_h_run = comm_h.add_run("Current CSV Comment:")
        comm_h_run.font.size = Pt(10)
        comm_h_run.font.bold = True
        comm_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
        
        comm_p = doc.add_paragraph()
        comm_run = comm_p.add_run(comments[:400])
        comm_run.font.size = Pt(9)
        comm_run.font.italic = True
        comm_p.paragraph_format.line_spacing = 1.15
        comm_p.paragraph_format.space_after = Pt(6)
    
    # After-fix checklist entry
    cklb_h = doc.add_paragraph()
    cklb_h_run = cklb_h.add_run("STIG Checklist Entry (after fix):")
    cklb_h_run.font.size = Pt(10)
    cklb_h_run.font.bold = True
    cklb_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
    
    cklb_table = doc.add_table(rows=3, cols=2)
    cklb_table.style = 'Table Grid'
    cklb_table.columns[0].width = Inches(2.0)
    cklb_table.columns[1].width = Inches(4.5)
    cklb_table.rows[0].cells[0].text = "Status"
    cklb_table.rows[0].cells[1].text = "Not a Finding"
    cklb_table.rows[1].cells[0].text = "Finding Details"
    cklb_table.rows[1].cells[1].text = f"Not a finding, the application is configured to meet the requirement per {admin_section}. See evidence screenshot for configuration verification."
    cklb_table.rows[2].cells[0].text = "Comments"
    cklb_table.rows[2].cells[1].text = f"See Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    for row in cklb_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        shade_cell(row.cells[0], "edf2f7")
    
    doc.add_page_break()

# Footer
last = doc.add_paragraph()
last_run = last.add_run("UNCLASSIFIED//CUI")
last_run.font.size = Pt(8)
last_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
last.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_dir = r"C:\Users\compj\.openclaw\workspace\output"
docx_path = os.path.join(output_dir, "Appian_ASD_STIG_V6R4_AdminConsole_Fixes.docx")
doc.save(docx_path)
print(f"DOCX saved: {docx_path}")

# Convert to PDF
import win32com.client as win32
word = win32.Dispatch('Word.Application')
word.Visible = False
pdf_path = docx_path.replace('.docx', '.pdf')
doc_w = word.Documents.Open(docx_path)
doc_w.SaveAs(pdf_path, FileFormat=17)
doc_w.Close()
word.Quit()
print(f"PDF saved: {pdf_path}")
