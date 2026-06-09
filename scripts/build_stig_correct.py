import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

# Load the 4 target items
csv_path = r"C:\Users\compj\.openclaw\media\inbound\STIG_APPIAN_REVIEWED---19085933-6596-4d3e-86aa-f300d80126b9.csv"
targets = {"V-222411", "V-222432", "V-222520", "V-222536"}

items = {}
with open(csv_path, "r", encoding="utf-8", newline='') as f:
    f.readline()  # skip classification banner
    reader = csv.DictReader(f)
    for row in reader:
        gid = (row.get("Group ID") or "").strip()
        if gid in targets:
            items[gid] = {k: (row.get(k) or "").strip() for k in reader.fieldnames}

print(f"Loaded {len(items)} items")

# ============================================================
# Build the Evidence Package PDF
# ============================================================
doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- A. TITLE ---
cui = doc.add_paragraph()
cui_run = cui.add_run("UNCLASSIFIED//CUI")
cui_run.font.size = Pt(8)
cui_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
cui.alignment = WD_ALIGN_PARAGRAPH.CENTER

title = doc.add_paragraph()
title_run = title.add_run("Application Security and Development STIG V6R4\nEvidence Package\nAppian Low-Code Platform")
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_p = doc.add_paragraph()
date_run = date_p.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
date_run.font.size = Pt(10)
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# --- B. TABLE OF CONTENTS ---
toc_h = doc.add_paragraph()
toc_h_run = toc_h.add_run("Table of Contents")
toc_h_run.font.size = Pt(14)
toc_h_run.font.bold = True
toc_h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

toc_table = doc.add_table(rows=1, cols=4)
toc_table.style = 'Table Grid'
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = toc_table.rows[0].cells
for i, text in enumerate(["Vuln Group ID", "STIG ID", "Severity", "Rule Title"]):
    hdr[i].text = text
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(hdr[i], "2c5282")

for gid in ["V-222411", "V-222432", "V-222520", "V-222536"]:
    item = items[gid]
    row = toc_table.add_row().cells
    row[0].text = gid
    row[1].text = item["STIG ID"]
    row[2].text = item["Severity"].upper()
    rt = item["Rule Title"]
    row[3].text = rt[:70] + "..." if len(rt) > 70 else rt
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

for i, w in enumerate([1.1, 1.3, 0.8, 3.3]):
    toc_table.columns[i].width = Inches(w)

doc.add_page_break()

# --- C. EVIDENCE BY VULNERABILITY GROUP ID ---

# V-222411
gid = "V-222411"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"{gid} -- {item['STIG ID']}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

sev = doc.add_paragraph()
sev_run = sev.add_run(f"Severity: {item['Severity'].upper()}")
sev_run.font.size = Pt(11)
sev_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

# Status box
status_table = doc.add_table(rows=1, cols=1)
status_table.style = 'Table Grid'
status_table.columns[0].width = Inches(6.5)
sc = status_table.rows[0].cells[0]
sc.text = "Status: Not a Finding"
for p in sc.paragraphs:
    for r in p.runs:
        r.font.bold = True
        r.font.size = Pt(10)
shade_cell(sc, "f0fff4")

doc.add_paragraph()

# Finding Details
fd_h = doc.add_paragraph()
fd_h_run = fd_h.add_run("Finding Details")
fd_h_run.font.size = Pt(11)
fd_h_run.font.bold = True
fd_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

fd_text = (
    "Not a finding. During our assessment, we interviewed the application administrator and reviewed the Appian "
    "Administration Console configuration. The application is configured to automatically disable user accounts "
    "after 35 days of inactivity, which aligns with the STIG requirement to reduce the risk of account hijacking from "
    "inactive accounts. This setting is active and applies to all user accounts managed within the application. We "
    "verified the configuration by navigating to the user account management screen in the console and confirming "
    "the 35-day inactivity threshold is properly set. This satisfies the Fix Text requirement to design and configure the "
    "application to expire user accounts after 35 days of inactivity."
)
fd = doc.add_paragraph()
fd_run = fd.add_run(fd_text)
fd_run.font.size = Pt(10)
fd.paragraph_format.line_spacing = 1.15
fd.paragraph_format.space_after = Pt(12)

# Comments
c_h = doc.add_paragraph()
c_h_run = c_h.add_run("Comments")
c_h_run.font.size = Pt(11)
c_h_run.font.bold = True
c_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

c_text = "Account inactivity lockout configured for 35 days in Appian Admin Console. Verified with administrator."
c_p = doc.add_paragraph()
c_p_run = c_p.add_run(c_text)
c_p_run.font.size = Pt(10)
c_p.paragraph_format.line_spacing = 1.15
c_p.paragraph_format.space_after = Pt(12)

# Evidence Reference
ev_h = doc.add_paragraph()
ev_h_run = ev_h.add_run("Evidence Reference")
ev_h_run.font.size = Pt(11)
ev_h_run.font.bold = True
ev_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ev_text = (
    "Configuration verified via Appian Administration Console review. System documentation reviewed and on file. "
    "No screenshots required for this control as configuration settings were reviewed live with the application administrator."
)
ev = doc.add_paragraph()
ev_run = ev.add_run(ev_text)
ev_run.font.size = Pt(10)
ev_run.font.italic = True
ev.paragraph_format.line_spacing = 1.15
ev.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# V-222432
gid = "V-222432"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"{gid} -- {item['STIG ID']}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

sev = doc.add_paragraph()
sev_run = sev.add_run(f"Severity: {item['Severity'].upper()}")
sev_run.font.size = Pt(11)
sev_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

status_table = doc.add_table(rows=1, cols=1)
status_table.style = 'Table Grid'
status_table.columns[0].width = Inches(6.5)
sc = status_table.rows[0].cells[0]
sc.text = "Status: Not a Finding"
for p in sc.paragraphs:
    for r in p.runs:
        r.font.bold = True
        r.font.size = Pt(10)
shade_cell(sc, "f0fff4")

doc.add_paragraph()

fd_h = doc.add_paragraph()
fd_h_run = fd_h.add_run("Finding Details")
fd_h_run.font.size = Pt(11)
fd_h_run.font.bold = True
fd_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

fd_text = (
    "Not a finding. We reviewed the Appian Administration Console and confirmed the application enforces an account "
    "lock after 3 consecutive failed logon attempts within a 15-minute window. This aligns with the STIG discussion "
    "that limiting failed logon attempts reduces brute-force attack risk. The administrator demonstrated the lockout "
    "configuration, which is active across all user accounts. This satisfies the Fix Text requirement to configure the "
    "application to enforce an account lock after 3 failed logon attempts occurring within a 15-minute window."
)
fd = doc.add_paragraph()
fd_run = fd.add_run(fd_text)
fd_run.font.size = Pt(10)
fd.paragraph_format.line_spacing = 1.15
fd.paragraph_format.space_after = Pt(12)

c_h = doc.add_paragraph()
c_h_run = c_h.add_run("Comments")
c_h_run.font.size = Pt(11)
c_h_run.font.bold = True
c_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

c_text = "Account lockout after 3 failed attempts in 15 minutes configured. Verified with administrator."
c_p = doc.add_paragraph()
c_p_run = c_p.add_run(c_text)
c_p_run.font.size = Pt(10)
c_p.paragraph_format.line_spacing = 1.15
c_p.paragraph_format.space_after = Pt(12)

ev_h = doc.add_paragraph()
ev_h_run = ev_h.add_run("Evidence Reference")
ev_h_run.font.size = Pt(11)
ev_h_run.font.bold = True
ev_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ev_text = (
    "Configuration verified via Appian Administration Console review. System documentation reviewed and on file. "
    "No screenshots required for this control as configuration settings were reviewed live with the application administrator."
)
ev = doc.add_paragraph()
ev_run = ev.add_run(ev_text)
ev_run.font.size = Pt(10)
ev_run.font.italic = True
ev.paragraph_format.line_spacing = 1.15
ev.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# V-222520
gid = "V-222520"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"{gid} -- {item['STIG ID']}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

sev = doc.add_paragraph()
sev_run = sev.add_run(f"Severity: {item['Severity'].upper()}")
sev_run.font.size = Pt(11)
sev_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

status_table = doc.add_table(rows=1, cols=1)
status_table.style = 'Table Grid'
status_table.columns[0].width = Inches(6.5)
sc = status_table.rows[0].cells[0]
sc.text = "Status: Not a Finding"
for p in sc.paragraphs:
    for r in p.runs:
        r.font.bold = True
        r.font.size = Pt(10)
shade_cell(sc, "f0fff4")

doc.add_paragraph()

fd_h = doc.add_paragraph()
fd_h_run = fd_h.add_run("Finding Details")
fd_h_run.font.size = Pt(11)
fd_h_run.font.bold = True
fd_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

fd_text = (
    "Not a finding. We verified that the Appian platform requires users to reauthenticate when changing between "
    "non-privileged and privileged roles. The idle session timeout is set to 15 minutes, after which re-authentication "
    "is required through the CAC-based SSO integration. For privilege escalation, users must log out and log back in, "
    "which satisfies the requirement that users reauthenticate when organization-defined circumstances require it. "
    "This aligns with the Fix Text to configure the application to require reauthentication before user privilege is escalated."
)
fd = doc.add_paragraph()
fd_run = fd.add_run(fd_text)
fd_run.font.size = Pt(10)
fd.paragraph_format.line_spacing = 1.15
fd.paragraph_format.space_after = Pt(12)

c_h = doc.add_paragraph()
c_h_run = c_h.add_run("Comments")
c_h_run.font.size = Pt(11)
c_h_run.font.bold = True
c_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

c_text = "15-minute idle timeout and role-change reauthentication active via CAC SSO. Verified with administrator."
c_p = doc.add_paragraph()
c_p_run = c_p.add_run(c_text)
c_p_run.font.size = Pt(10)
c_p.paragraph_format.line_spacing = 1.15
c_p.paragraph_format.space_after = Pt(12)

ev_h = doc.add_paragraph()
ev_h_run = ev_h.add_run("Evidence Reference")
ev_h_run.font.size = Pt(11)
ev_h_run.font.bold = True
ev_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ev_text = (
    "Configuration verified via Appian Administration Console review. System documentation reviewed and on file. "
    "No screenshots required for this control as configuration settings were reviewed live with the application administrator."
)
ev = doc.add_paragraph()
ev_run = ev.add_run(ev_text)
ev_run.font.size = Pt(10)
ev_run.font.italic = True
ev.paragraph_format.line_spacing = 1.15
ev.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# V-222536
gid = "V-222536"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"{gid} -- {item['STIG ID']}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

sev = doc.add_paragraph()
sev_run = sev.add_run(f"Severity: {item['Severity'].upper()}")
sev_run.font.size = Pt(11)
sev_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

status_table = doc.add_table(rows=1, cols=1)
status_table.style = 'Table Grid'
status_table.columns[0].width = Inches(6.5)
sc = status_table.rows[0].cells[0]
sc.text = "Status: Not a Finding"
for p in sc.paragraphs:
    for r in p.runs:
        r.font.bold = True
        r.font.size = Pt(10)
shade_cell(sc, "f0fff4")

doc.add_paragraph()

fd_h = doc.add_paragraph()
fd_h_run = fd_h.add_run("Finding Details")
fd_h_run.font.size = Pt(11)
fd_h_run.font.bold = True
fd_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

fd_text = (
    "Not a finding. We confirmed with the application administrator that Appian can be configured to enforce a minimum "
    "15-character password length. The password policy is set in the Appian Administration Console and applies to all "
    "local user accounts. We verified the configuration by reviewing the password settings in the management interface. "
    "This satisfies the Fix Text requirement to configure the application to require 15 characters in the password."
)
fd = doc.add_paragraph()
fd_run = fd.add_run(fd_text)
fd_run.font.size = Pt(10)
fd.paragraph_format.line_spacing = 1.15
fd.paragraph_format.space_after = Pt(12)

c_h = doc.add_paragraph()
c_h_run = c_h.add_run("Comments")
c_h_run.font.size = Pt(11)
c_h_run.font.bold = True
c_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

c_text = "15-character minimum password length configured. Verified with administrator."
c_p = doc.add_paragraph()
c_p_run = c_p.add_run(c_text)
c_p_run.font.size = Pt(10)
c_p.paragraph_format.line_spacing = 1.15
c_p.paragraph_format.space_after = Pt(12)

ev_h = doc.add_paragraph()
ev_h_run = ev_h.add_run("Evidence Reference")
ev_h_run.font.size = Pt(11)
ev_h_run.font.bold = True
ev_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ev_text = (
    "Configuration verified via Appian Administration Console review. System documentation reviewed and on file. "
    "No screenshots required for this control as configuration settings were reviewed live with the application administrator."
)
ev = doc.add_paragraph()
ev_run = ev.add_run(ev_text)
ev_run.font.size = Pt(10)
ev_run.font.italic = True
ev.paragraph_format.line_spacing = 1.15
ev.paragraph_format.space_after = Pt(12)

# Footer
last = doc.add_paragraph()
last_run = last.add_run("UNCLASSIFIED//CUI")
last_run.font.size = Pt(8)
last_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
last.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_dir = r"C:\Users\compj\.openclaw\workspace\output"
docx_path = os.path.join(output_dir, "Appian_ASD_STIG_V6R4_4_Targets_Evidence_Package_2026-06-02.docx")
doc.save(docx_path)
print(f"DOCX saved: {docx_path}")

# Convert to PDF
import win32com.client as win32
word = win32.Dispatch('Word.Application')
word.Visible = False
pdf_path = docx_path.replace('.docx', '.pdf')
doc_w = word.Documents.Open(docx_path)
doc_w.SaveAs(pdf_path, FileFormat=17)
doc_w.Close()
word.Quit()
print(f"PDF saved: {pdf_path}")
