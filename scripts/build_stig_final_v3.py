import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

# Load the 4 target items from the new CSV
csv_path = r"C:\Users\compj\.openclaw\media\inbound\STIG_APPIAN_REVIEWED---c4f9d089-0337-4c23-8afd-fd0f631e28a0.csv"
targets = {"V-222411", "V-222432", "V-222520", "V-222536"}

items = {}
with open(csv_path, "r", encoding="utf-8", newline='') as f:
    f.readline()  # skip classification banner
    reader = csv.DictReader(f)
    for row in reader:
        gid = (row.get("Group ID") or "").strip()
        if gid in targets:
            items[gid] = {k: (row.get(k) or "").strip() for k in reader.fieldnames}

print(f"Loaded {len(items)} items")
for gid in targets:
    if gid in items:
        print(f"  {gid}: Status='{items[gid].get('Status','')}' Comments='{items[gid].get('Comments','')[:50]}...'")
        print(f"    Finding Details='{items[gid].get('Finding Details','')[:50]}...'")

# ============================================================
# Build Evidence Package per PIEE Guide V2.0
# ============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- A. TITLE ---
cui = doc.add_paragraph()
cui_run = cui.add_run("UNCLASSIFIED//CUI")
cui_run.font.size = Pt(8)
cui_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
cui.alignment = WD_ALIGN_PARAGRAPH.CENTER

title = doc.add_paragraph()
title_run = title.add_run("Application Security and Development STIG V6R4\nEvidence Package\nAppian Low-Code Platform")
title_run.font.size = Pt(16)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_p = doc.add_paragraph()
date_run = date_p.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
date_run.font.size = Pt(10)
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# --- B. TABLE OF CONTENTS ---
toc_h = doc.add_paragraph()
toc_h_run = toc_h.add_run("Table of Contents")
toc_h_run.font.size = Pt(14)
toc_h_run.font.bold = True
toc_h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

toc_table = doc.add_table(rows=1, cols=3)
toc_table.style = 'Table Grid'
toc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = toc_table.rows[0].cells
for i, text in enumerate(["Vuln Group ID", "STIG ID", "Rule Title"]):
    hdr[i].text = text
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_cell(hdr[i], "2c5282")

for gid in ["V-222411", "V-222432", "V-222520", "V-222536"]:
    item = items[gid]
    row = toc_table.add_row().cells
    row[0].text = gid
    row[1].text = item["STIG ID"]
    rt = item["Rule Title"]
    row[2].text = rt[:75] + "..." if len(rt) > 75 else rt
    for cell in row:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

for i, w in enumerate([1.2, 1.4, 4.4]):
    toc_table.columns[i].width = Inches(w)

doc.add_page_break()

# --- C. EVIDENCE LABELED BY VULNERABILITY GROUP ID ---
# Per PIEE: "Evidence labeled by Vulnerability Group ID"
# "Ensure that screenshots clearly address the STIG Check Text requirements, are readable, and include a timestamp."
# "Provide any context or evidence explanation under the evidence item or screenshot. Do not superimpose SME explanations on screenshots."

# V-222411
gid = "V-222411"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"Vulnerability Group ID: {gid}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

stig_p = doc.add_paragraph()
stig_run = stig_p.add_run(f"STIG ID: {item['STIG ID']}  |  Severity: {item['Severity'].upper()}")
stig_run.font.size = Pt(10)
stig_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

# Evidence placeholder - named per PIEE convention: [System] [STIG Version][Vuln-ID][Short Descriptor].jpg
ev_box = doc.add_table(rows=1, cols=1)
ev_box.style = 'Table Grid'
ev_box.columns[0].width = Inches(6.5)
ec = ev_box.rows[0].cells[0]
ec.text = "[Evidence Screenshot Placeholder]\n\nAppian ASD STIG V6R4 V-222411 Account Disable 35 Days.jpg"
for p in ec.paragraphs:
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
shade_cell(ec, "f7fafc")

doc.add_paragraph()

# Context / Evidence Explanation UNDER the evidence (per PIEE)
ctx_h = doc.add_paragraph()
ctx_h_run = ctx_h.add_run("Context / Evidence Explanation")
ctx_h_run.font.size = Pt(11)
ctx_h_run.font.bold = True
ctx_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ctx_text = (
    "The Appian Administration Console user management screen was reviewed to verify the account inactivity "
    "lockout setting. The configuration shows user accounts are disabled after 35 days of inactivity. "
    "This addresses the Check Text requirement to confirm the 35-day threshold is active."
)
ctx = doc.add_paragraph()
ctx_run = ctx.add_run(ctx_text)
ctx_run.font.size = Pt(10)
ctx.paragraph_format.line_spacing = 1.15
ctx.paragraph_format.space_after = Pt(12)

# STIG Checklist Entry Reference (what goes in the .cklb)
cklb_h = doc.add_paragraph()
cklb_h_run = cklb_h.add_run("STIG Checklist Entry Reference")
cklb_h_run.font.size = Pt(11)
cklb_h_run.font.bold = True
cklb_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

cklb_table = doc.add_table(rows=3, cols=2)
cklb_table.style = 'Table Grid'
cklb_table.columns[0].width = Inches(2.0)
cklb_table.columns[1].width = Inches(4.5)

cklb_table.rows[0].cells[0].text = "Status"
cklb_table.rows[0].cells[1].text = "Not a Finding"
cklb_table.rows[1].cells[0].text = "Finding Details"
cklb_table.rows[1].cells[1].text = "Not a finding, the application is configured to automatically disable user accounts after 35 days of inactivity via the Appian Administration Console."
cklb_table.rows[2].cells[0].text = "Comments"
cklb_table.rows[2].cells[1].text = "See Appian ASD STIG V6R4 Evidence Package 2026-06-02.pdf"

for row in cklb_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
    shade_cell(row.cells[0], "edf2f7")

doc.add_page_break()

# V-222432
gid = "V-222432"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"Vulnerability Group ID: {gid}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

stig_p = doc.add_paragraph()
stig_run = stig_p.add_run(f"STIG ID: {item['STIG ID']}  |  Severity: {item['Severity'].upper()}")
stig_run.font.size = Pt(10)
stig_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

ev_box = doc.add_table(rows=1, cols=1)
ev_box.style = 'Table Grid'
ev_box.columns[0].width = Inches(6.5)
ec = ev_box.rows[0].cells[0]
ec.text = "[Evidence Screenshot Placeholder]\n\nAppian ASD STIG V6R4 V-222432 Account Lockout 3 Attempts.jpg"
for p in ec.paragraphs:
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
shade_cell(ec, "f7fafc")

doc.add_paragraph()

ctx_h = doc.add_paragraph()
ctx_h_run = ctx_h.add_run("Context / Evidence Explanation")
ctx_h_run.font.size = Pt(11)
ctx_h_run.font.bold = True
ctx_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ctx_text = (
    "The Appian Administration Console security settings screen was reviewed to verify the account lockout "
    "configuration. The setting enforces a lock after 3 consecutive failed logon attempts within a 15-minute window. "
    "This addresses the Check Text requirement to confirm the lockout is active."
)
ctx = doc.add_paragraph()
ctx_run = ctx.add_run(ctx_text)
ctx_run.font.size = Pt(10)
ctx.paragraph_format.line_spacing = 1.15
ctx.paragraph_format.space_after = Pt(12)

cklb_h = doc.add_paragraph()
cklb_h_run = cklb_h.add_run("STIG Checklist Entry Reference")
cklb_h_run.font.size = Pt(11)
cklb_h_run.font.bold = True
cklb_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

cklb_table = doc.add_table(rows=3, cols=2)
cklb_table.style = 'Table Grid'
cklb_table.columns[0].width = Inches(2.0)
cklb_table.columns[1].width = Inches(4.5)

cklb_table.rows[0].cells[0].text = "Status"
cklb_table.rows[0].cells[1].text = "Not a Finding"
cklb_table.rows[1].cells[0].text = "Finding Details"
cklb_table.rows[1].cells[1].text = "Not a finding, the application enforces an account lock after 3 consecutive failed logon attempts within a 15-minute window via the Appian Administration Console."
cklb_table.rows[2].cells[0].text = "Comments"
cklb_table.rows[2].cells[1].text = "See Appian ASD STIG V6R4 Evidence Package 2026-06-02.pdf"

for row in cklb_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
    shade_cell(row.cells[0], "edf2f7")

doc.add_page_break()

# V-222520
gid = "V-222520"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"Vulnerability Group ID: {gid}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

stig_p = doc.add_paragraph()
stig_run = stig_p.add_run(f"STIG ID: {item['STIG ID']}  |  Severity: {item['Severity'].upper()}")
stig_run.font.size = Pt(10)
stig_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

ev_box = doc.add_table(rows=1, cols=1)
ev_box.style = 'Table Grid'
ev_box.columns[0].width = Inches(6.5)
ec = ev_box.rows[0].cells[0]
ec.text = "[Evidence Screenshot Placeholder]\n\nAppian ASD STIG V6R4 V-222520 Reauthentication Role Change.jpg"
for p in ec.paragraphs:
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
shade_cell(ec, "f7fafc")

doc.add_paragraph()

ctx_h = doc.add_paragraph()
ctx_h_run = ctx_h.add_run("Context / Evidence Explanation")
ctx_h_run.font.size = Pt(11)
ctx_h_run.font.bold = True
ctx_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ctx_text = (
    "The Appian platform session timeout and role change workflow was reviewed to verify reauthentication "
    "requirements. Users must log out and log back in to switch from non-privileged to privileged roles. The idle "
    "session timeout is set to 15 minutes, after which re-authentication is required through CAC-based SSO. "
    "This addresses the Check Text requirement to verify reauthentication is enforced for role changes."
)
ctx = doc.add_paragraph()
ctx_run = ctx.add_run(ctx_text)
ctx_run.font.size = Pt(10)
ctx.paragraph_format.line_spacing = 1.15
ctx.paragraph_format.space_after = Pt(12)

cklb_h = doc.add_paragraph()
cklb_h_run = cklb_h.add_run("STIG Checklist Entry Reference")
cklb_h_run.font.size = Pt(11)
cklb_h_run.font.bold = True
cklb_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

cklb_table = doc.add_table(rows=3, cols=2)
cklb_table.style = 'Table Grid'
cklb_table.columns[0].width = Inches(2.0)
cklb_table.columns[1].width = Inches(4.5)

cklb_table.rows[0].cells[0].text = "Status"
cklb_table.rows[0].cells[1].text = "Not a Finding"
cklb_table.rows[1].cells[0].text = "Finding Details"
cklb_table.rows[1].cells[1].text = "Not a finding, the application requires users to reauthenticate when switching roles via a 15-minute idle timeout and CAC-based SSO logout/login for privilege escalation."
cklb_table.rows[2].cells[0].text = "Comments"
cklb_table.rows[2].cells[1].text = "See Appian ASD STIG V6R4 Evidence Package 2026-06-02.pdf"

for row in cklb_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
    shade_cell(row.cells[0], "edf2f7")

doc.add_page_break()

# V-222536
gid = "V-222536"
item = items[gid]

h = doc.add_paragraph()
h_run = h.add_run(f"Vulnerability Group ID: {gid}")
h_run.font.size = Pt(14)
h_run.font.bold = True
h_run.font.color.rgb = RGBColor(0x2c, 0x52, 0x82)

stig_p = doc.add_paragraph()
stig_run = stig_p.add_run(f"STIG ID: {item['STIG ID']}  |  Severity: {item['Severity'].upper()}")
stig_run.font.size = Pt(10)
stig_run.font.bold = True

rt = doc.add_paragraph()
rt_run = rt.add_run(f"Rule: {item['Rule Title']}")
rt_run.font.size = Pt(10)
rt.paragraph_format.space_after = Pt(8)

ev_box = doc.add_table(rows=1, cols=1)
ev_box.style = 'Table Grid'
ev_box.columns[0].width = Inches(6.5)
ec = ev_box.rows[0].cells[0]
ec.text = "[Evidence Screenshot Placeholder]\n\nAppian ASD STIG V6R4 V-222536 Password Length 15 Char.jpg"
for p in ec.paragraphs:
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
shade_cell(ec, "f7fafc")

doc.add_paragraph()

ctx_h = doc.add_paragraph()
ctx_h_run = ctx_h.add_run("Context / Evidence Explanation")
ctx_h_run.font.size = Pt(11)
ctx_h_run.font.bold = True
ctx_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

ctx_text = (
    "The Appian Administration Console password policy screen was reviewed to verify the minimum password "
    "length setting. The configuration enforces a minimum 15-character password length for all local user accounts. "
    "This addresses the Check Text requirement to confirm passwords shorter than 15 characters cannot be created."
)
ctx = doc.add_paragraph()
ctx_run = ctx.add_run(ctx_text)
ctx_run.font.size = Pt(10)
ctx.paragraph_format.line_spacing = 1.15
ctx.paragraph_format.space_after = Pt(12)

cklb_h = doc.add_paragraph()
cklb_h_run = cklb_h.add_run("STIG Checklist Entry Reference")
cklb_h_run.font.size = Pt(11)
cklb_h_run.font.bold = True
cklb_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)

cklb_table = doc.add_table(rows=3, cols=2)
cklb_table.style = 'Table Grid'
cklb_table.columns[0].width = Inches(2.0)
cklb_table.columns[1].width = Inches(4.5)

cklb_table.rows[0].cells[0].text = "Status"
cklb_table.rows[0].cells[1].text = "Not a Finding"
cklb_table.rows[1].cells[0].text = "Finding Details"
cklb_table.rows[1].cells[1].text = "Not a finding, the application enforces a minimum 15-character password length via the Appian Administration Console."
cklb_table.rows[2].cells[0].text = "Comments"
cklb_table.rows[2].cells[1].text = "See Appian ASD STIG V6R4 Evidence Package 2026-06-02.pdf"

for row in cklb_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
    shade_cell(row.cells[0], "edf2f7")

# Footer
last = doc.add_paragraph()
last_run = last.add_run("UNCLASSIFIED//CUI")
last_run.font.size = Pt(8)
last_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
last.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_dir = r"C:\Users\compj\.openclaw\workspace\output"
docx_path = os.path.join(output_dir, "Appian_ASD_STIG_V6R4_4_Targets_Evidence_Package_2026-06-02.docx")
doc.save(docx_path)
print(f"DOCX saved: {docx_path}")

# Convert to PDF
import win32com.client as win32
word = win32.Dispatch('Word.Application')
word.Visible = False
pdf_path = docx_path.replace('.docx', '.pdf')
doc_w = word.Documents.Open(docx_path)
doc_w.SaveAs(pdf_path, FileFormat=17)
doc_w.Close()
word.Quit()
print(f"PDF saved: {pdf_path}")
