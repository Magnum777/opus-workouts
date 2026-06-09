import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade_cell(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

# Load items from CSV
csv_path = r"C:\Users\compj\.openclaw\media\inbound\STIG_APPIAN_REVIEWED---c4f9d089-0337-4c23-8afd-fd0f631e28a0.csv"
all_items = {}
with open(csv_path, "r", encoding="utf-8", newline='') as f:
    f.readline()
    reader = csv.DictReader(f)
    for row in reader:
        gid = (row.get("Group ID") or "").strip()
        all_items[gid] = row

# Only the 4 verified Not a Finding items
target_ids = ["V-222411", "V-222432", "V-222520", "V-222536"]

# ============================================================
# Build lean evidence package per PIEE template
# ============================================================
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# --- A. TITLE ---
title = doc.add_paragraph()
title_run = title.add_run("Application Security and Development STIG\nVersion: 6, Release: 4")
title_run.font.size = Pt(20)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

cui = doc.add_paragraph()
cui_run = cui.add_run("UNCLASSIFIED//CUI")
cui_run.font.size = Pt(10)
cui_run.font.bold = True

doc.add_paragraph()

# --- B. CONTENTS ---
contents_h = doc.add_paragraph()
contents_h_run = contents_h.add_run("Contents")
contents_h_run.font.size = Pt(14)
contents_h_run.font.bold = True
contents_h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

page_num = 3
for gid in target_ids:
    toc_line = doc.add_paragraph()
    toc_run = toc_line.add_run(f"{gid} (Not a Finding)")
    toc_run.font.size = Pt(11)
    dots = " ." * (70 - len(gid) - 17)
    dots_run = toc_line.add_run(f"{dots} {page_num}")
    dots_run.font.size = Pt(11)
    toc_line.paragraph_format.space_after = Pt(2)
    page_num += 1

doc.add_page_break()

# ============================================================
# EVIDENCE BY VULNERABILITY GROUP ID
# ============================================================
short_desc_map = {
    "V-222411": "Account Disable 35 Days",
    "V-222432": "Account Lockout 3 Attempts",
    "V-222520": "Reauthentication Role Change",
    "V-222536": "Password Length 15 Char"
}

explanations = {
    "V-222411": (
        "The Appian Administration Console user management screen was reviewed to verify the account "
        "inactivity lockout setting. The configuration shows user accounts are disabled after 35 days "
        "of inactivity. This addresses the Check Text requirement to confirm the 35-day threshold is active."
    ),
    "V-222432": (
        "The Appian Administration Console security settings screen was reviewed to verify the account "
        "lockout configuration. The setting enforces a lock after 3 consecutive failed logon attempts "
        "within a 15-minute window. This addresses the Check Text requirement to confirm the lockout is active."
    ),
    "V-222520": (
        "The Appian platform session timeout and role change workflow was reviewed to verify reauthentication "
        "requirements. Users must log out and log back in to switch from non-privileged to privileged roles. "
        "The idle session timeout is set to 15 minutes, after which re-authentication is required through CAC-based SSO. "
        "This addresses the Check Text requirement to verify reauthentication is enforced for role changes."
    ),
    "V-222536": (
        "The Appian Administration Console password policy screen was reviewed to verify the minimum "
        "password length setting. The configuration enforces a minimum 15-character password length "
        "for all local user accounts. This addresses the Check Text requirement to confirm passwords "
        "shorter than 15 characters cannot be created."
    )
}

for gid in target_ids:
    item = all_items[gid]
    short_desc = short_desc_map[gid]
    
    # Vuln ID header in blue
    h = doc.add_paragraph()
    h_run = h.add_run(gid)
    h_run.font.size = Pt(14)
    h_run.font.bold = True
    h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    h.paragraph_format.space_after = Pt(6)
    
    # Evidence placeholder (screenshot)
    ev_box = doc.add_table(rows=1, cols=1)
    ev_box.style = 'Table Grid'
    ev_box.columns[0].width = Inches(6.5)
    ec = ev_box.rows[0].cells[0]
    ev_filename = f"Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    ec.text = f"[Evidence Screenshot Placeholder]\n\n{ev_filename}"
    for p in ec.paragraphs:
        for r in p.runs:
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    shade_cell(ec, "f5f5f5")
    
    doc.add_paragraph()
    
    # Explanation directly under evidence (per PIEE example)
    exp_p = doc.add_paragraph()
    exp_label = exp_p.add_run("Explanation/Context: ")
    exp_label.font.size = Pt(10)
    exp_label.font.bold = True
    exp_text = exp_p.add_run(f"({explanations[gid]})")
    exp_text.font.size = Pt(10)
    exp_p.paragraph_format.line_spacing = 1.2
    exp_p.paragraph_format.space_after = Pt(12)
    
    # STIG Checklist Entry Reference
    cklb_h = doc.add_paragraph()
    cklb_h_run = cklb_h.add_run("STIG Checklist Entry Reference")
    cklb_h_run.font.size = Pt(10)
    cklb_h_run.font.bold = True
    cklb_h_run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
    
    cklb_table = doc.add_table(rows=3, cols=2)
    cklb_table.style = 'Table Grid'
    cklb_table.columns[0].width = Inches(2.0)
    cklb_table.columns[1].width = Inches(4.5)
    cklb_table.rows[0].cells[0].text = "Status"
    cklb_table.rows[0].cells[1].text = "Not a Finding"
    cklb_table.rows[1].cells[0].text = "Finding Details"
    fd = item.get('Finding Details', '').strip()
    if not fd:
        fd = f"Not a finding, the application is configured to meet the requirement per {short_desc}."
    cklb_table.rows[1].cells[1].text = fd[:300]
    cklb_table.rows[2].cells[0].text = "Comments"
    cklb_table.rows[2].cells[1].text = f"See Appian ASD STIG V6R4 {gid} {short_desc}.jpg"
    
    for row in cklb_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
        shade_cell(row.cells[0], "edf2f7")
    
    doc.add_page_break()

# --- Supplemental Standalone Evidence ---
supp_h = doc.add_paragraph()
supp_h_run = supp_h.add_run("Supplemental Standalone Evidence")
supp_h_run.font.size = Pt(12)
supp_h_run.font.bold = True
supp_h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

bullet1 = doc.add_paragraph(style='List Bullet')
b1_run = bullet1.add_run(
    "If evidence item is an organization-level or vendor document then please also submit the whole "
    "document separately.\n"
    "Example: PIEE COE Access Control Policy 07012024.pdf"
)
b1_run.font.size = Pt(10)

bullet2 = doc.add_paragraph(style='List Bullet')
b2_run = bullet2.add_run(
    "If submitting a mixed STIG (Automated + Manual input), you can submit standalone files named "
    "in accordance with our convention [System/Application] [STIG and Version/Revision][Vuln-ID]"
    "[Short Descriptor].jpg.\n"
    "Example: PIEE eBusiness Suite ASD STIG V5R3 V-222645 Integrity Mismatch Check 062024.jpg"
)
b2_run.font.size = Pt(10)

doc.add_paragraph()

# --- 5. CUI Markings ---
cui_h = doc.add_paragraph()
cui_h_run = cui_h.add_run("5. CUI Markings")
cui_h_run.font.size = Pt(12)
cui_h_run.font.bold = True
cui_h_run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

cui_text = doc.add_paragraph()
cui_text_run = cui_text.add_run(
    "The SME POC is responsible for ensuring all submissions are properly marked before submitting to the PIEE PMO. "
    "This may include but is not limited to:\n\n"
    "UNCLASSIFIED//CUI"
)
cui_text_run.font.size = Pt(10)

# Footer
last = doc.add_paragraph()
last_run = last.add_run("UNCLASSIFIED//CUI")
last_run.font.size = Pt(8)
last_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
last.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_dir = r"C:\Users\compj\.openclaw\workspace\output"
docx_path = os.path.join(output_dir, "Appian_ASD_STIG_V6R4_Evidence_Package.docx")
doc.save(docx_path)
print(f"DOCX saved: {docx_path}")

# Convert to PDF
import win32com.client as win32
word = win32.Dispatch('Word.Application')
word.Visible = False
pdf_path = docx_path.replace('.docx', '.pdf')
doc_w = word.Documents.Open(docx_path)
doc_w.SaveAs(pdf_path, FileFormat=17)
doc_w.Close()
word.Quit()
print(f"PDF saved: {pdf_path}")
